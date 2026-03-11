# -*- coding: utf-8 -*-
"""
1. 替换图4-1图片（图内已无标题，不需要删图题，但要删掉图片前多余的标题段落）
2. 替换图4-7a、图4-7b图片
"""
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
DIAG = r'd:\SoftWare\Code\bysj\diagrams'

def log(s):
    sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def has_img(para):
    return para._element.find('.//' + qn('w:drawing')) is not None

def find_idx(*keys, start=0):
    for i in range(start, len(doc.paragraphs)):
        if all(k in doc.paragraphs[i].text for k in keys):
            return i
    return -1

def make_img_elem(img_path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Cm(width_cm))
    e = p._element
    e.getparent().remove(e)
    return e

def replace_img_para(img_idx, new_img_path, width_cm):
    """用新图片段落替换指定索引的图片段落"""
    old_elem = doc.paragraphs[img_idx]._element
    parent = old_elem.getparent()
    pos = list(parent).index(old_elem)
    new_elem = make_img_elem(new_img_path, width_cm)
    parent.insert(pos, new_elem)
    parent.remove(old_elem)
    log(f'  替换图片段落 [{img_idx}]')

def remove_para_at(idx):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)
    log(f'  删除段落 [{idx}]: {doc.paragraphs[idx].text[:30] if idx < len(doc.paragraphs) else ""}')

# ── 诊断：列出涉及图的段落 ─────────────────────────────────
log('=== 诊断：图片与图题位置 ===')
for i in range(150, 210):
    p = doc.paragraphs[i]
    t = p.text.strip()
    img = has_img(p)
    if img or ('图 4-' in t) or ('图4-' in t):
        log(f'  [{i}] img={img}  "{t[:60]}"')

# ══════════════════════════════════════════════════════════
# 1. 图4-1：找图片段落和图题段落
# ══════════════════════════════════════════════════════════
log('\n--- 处理图4-1 ---')
cap41 = find_idx('图 4-1', '系统总体架构图')
log(f'  图题段落: [{cap41}] "{doc.paragraphs[cap41].text if cap41>0 else ""}"')

# 在图题前找图片段落
img41 = -1
for i in range(cap41-1, max(cap41-6,-1), -1):
    if has_img(doc.paragraphs[i]):
        img41 = i; break
log(f'  图片段落: [{img41}]')

# 检查图片前是否还有一个多余的"图4-1..."标题段落
if img41 > 0:
    for i in range(img41-1, max(img41-4,-1), -1):
        t = doc.paragraphs[i].text.strip()
        if '图 4-1' in t or '图4-1' in t:
            log(f'  发现多余标题段落 [{i}]: "{t}"')
            remove_para_at(i)
            # 索引偏移，重新找
            img41 = find_idx('图 4-1', '系统总体架构图')
            cap41 = img41  # 重新用find找图片
            for j in range(cap41-1, max(cap41-6,-1), -1):
                if has_img(doc.paragraphs[j]):
                    img41 = j; break
            break

# 替换图片
if img41 >= 0:
    replace_img_para(img41, os.path.join(DIAG, 'fig4-1_arch.png'), 13.5)

# ══════════════════════════════════════════════════════════
# 2. 图4-7a
# ══════════════════════════════════════════════════════════
log('\n--- 处理图4-7a ---')
cap7a = find_idx('图 4-7a')
log(f'  图4-7a图题: [{cap7a}]')
img7a = -1
for i in range(cap7a-1, max(cap7a-6,-1), -1):
    if has_img(doc.paragraphs[i]):
        img7a = i; break
# 也可能在图题之后
if img7a < 0:
    for i in range(cap7a+1, min(cap7a+6, len(doc.paragraphs))):
        if has_img(doc.paragraphs[i]):
            img7a = i; break
log(f'  图4-7a图片: [{img7a}]')
if img7a >= 0:
    replace_img_para(img7a, os.path.join(DIAG, 'fig4-7a_er_core.png'), 13.5)

# ══════════════════════════════════════════════════════════
# 3. 图4-7b
# ══════════════════════════════════════════════════════════
log('\n--- 处理图4-7b ---')
cap7b = find_idx('图 4-7b')
log(f'  图4-7b图题: [{cap7b}]')
img7b = -1
for i in range(cap7b-1, max(cap7b-6,-1), -1):
    if has_img(doc.paragraphs[i]):
        img7b = i; break
if img7b < 0:
    for i in range(cap7b+1, min(cap7b+6, len(doc.paragraphs))):
        if has_img(doc.paragraphs[i]):
            img7b = i; break
log(f'  图4-7b图片: [{img7b}]')
if img7b >= 0:
    replace_img_para(img7b, os.path.join(DIAG, 'fig4-7b_er_msg.png'), 13.5)

# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存。')
