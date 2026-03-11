# -*- coding: utf-8 -*-
"""
1. 重写 4.1 正文两段（[152][153]）
2. 替换图 4-1 图片
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, os, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
DIAG = r'd:\SoftWare\Code\bysj\diagrams'

def log(s):
    sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

# ─────────────────────────────────────────────────────────
# 重写段落文字（清空runs，写新文字，取消加粗）
# ─────────────────────────────────────────────────────────
def set_text(idx, new_text):
    para = doc.paragraphs[idx]
    for run in para.runs:
        run.text = ''
        run.font.bold = False
    if para.runs:
        para.runs[0].text = new_text
    else:
        r = para.add_run(new_text)
        r.font.bold = False

# ─────────────────────────────────────────────────────────
# 1. 重写 4.1 正文
# ─────────────────────────────────────────────────────────
log('--- 更新 4.1 正文 ---')

text_152 = (
    '本系统采用成熟的 B/S（Browser/Server）架构，用户通过浏览器发起请求，'
    '后端服务器负责业务处理与数据管理，实现客户端与服务器的职责分离。'
    '前端基于 Bootstrap 5 框架进行响应式页面设计，可自适应 PC、平板及手机等多终端访问，'
    '无需额外安装客户端软件。后端采用 Django 框架提供的 MVT（Model-View-Template）设计模式，'
    '将数据建模、业务逻辑与页面渲染三个职责明确分离，便于模块化开发与后期维护。'
    '系统总体架构如图 4-1 所示。'
)

text_153 = (
    '在 MVT 架构中，模型层（Model）通过 Django ORM 与 MySQL 数据库进行映射，'
    '负责定义用户、家庭药箱、共享帖子等核心实体及其关联关系，并承担数据校验与持久化操作；'
    '视图层（View）承担系统的核心业务逻辑，包括多级审批流的状态流转、'
    '基于角色的访问控制（RBAC）权限校验以及系统通知的聚合与分发；'
    '模板层（Template）负责将视图层传递的上下文数据渲染为 HTML 页面，'
    '并根据当前登录用户的角色动态控制操作菜单的显示内容，实现前端界面与后端数据的解耦。'
)

set_text(152, text_152)
set_text(153, text_153)
log('  4.1 正文已更新')

# ─────────────────────────────────────────────────────────
# 2. 替换图 4-1 图片
# ─────────────────────────────────────────────────────────
log('--- 替换图 4-1 ---')

# 找到图片段落 [154] 和图题 [155]
img_idx = 154
cap_idx = 155

# 验证
has_img = doc.paragraphs[img_idx]._element.find('.//' + qn('w:drawing')) is not None
log(f'  [154] has_img={has_img}  [155] text={doc.paragraphs[cap_idx].text[:30]}')

# 创建新图片段落
p_new = doc.add_paragraph()
p_new.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_new.add_run()
run.add_picture(os.path.join(DIAG, 'fig4-1_arch.png'), width=Cm(13.5))
new_img_elem = p_new._element
new_img_elem.getparent().remove(new_img_elem)

# 替换旧图片段落
old_img_elem = doc.paragraphs[img_idx]._element
parent = old_img_elem.getparent()
pos = list(parent).index(old_img_elem)
parent.insert(pos, new_img_elem)
parent.remove(old_img_elem)

log('  图 4-1 替换完成')

# ─────────────────────────────────────────────────────────
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存。')
