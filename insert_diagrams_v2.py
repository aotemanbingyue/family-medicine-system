# -*- coding: utf-8 -*-
"""
正确插入所有流程图和ER图到Word文档对应位置
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, os, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
DIAG = r'd:\SoftWare\Code\bysj\diagrams'

def log(s):
    sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def find_idx(*keys, start=0, end=None):
    end = end or len(doc.paragraphs)
    for i in range(start, end):
        if all(k in doc.paragraphs[i].text for k in keys):
            return i
    return -1

def make_img_elem(img_path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Cm(width_cm))
    e = p._element
    e.getparent().remove(e)
    return e

def make_caption_elem(text, ref_para_idx):
    ref = doc.paragraphs[ref_para_idx]
    new_p = OxmlElement('w:p')
    ref_pPr = ref._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); new_p.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(qn('w:val'), 'center')
    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr.append(OxmlElement('w:b'))
    new_r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(t)
    new_p.append(new_r)
    return new_p

def insert_before(anchor_idx, *elems):
    """在 anchor_idx 段落前插入（注意：每次插入后anchor_idx后移）"""
    anchor = doc.paragraphs[anchor_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)
    for elem in elems:
        parent.insert(pos, elem)
        pos += 1

# ══════════════════════════════════════════════════════════
# 1. 图4-2 替换：找到[160]"图 4-2  系统功能模块图"前的图片，替换为新图
# ══════════════════════════════════════════════════════════
log('--- 处理图4-2 ---')
# 找到图题
cap42_idx = find_idx('图 4-2', '系统功能模块图')
log(f'  图4-2图题位置: [{cap42_idx}]')

if cap42_idx > 0:
    # 找到前面的图片段落
    from docx.oxml.ns import qn as _qn
    img42_idx = -1
    for i in range(cap42_idx - 1, max(cap42_idx - 5, -1), -1):
        if doc.paragraphs[i]._element.find('.//' + _qn('w:drawing')) is not None:
            img42_idx = i
            break
    log(f'  旧图片位置: [{img42_idx}]')

    # 在图题前插入新图片
    new_img = make_img_elem(os.path.join(DIAG, 'fig4-2_module_tree.png'), 14.0)
    cap_elem = doc.paragraphs[cap42_idx]._element
    parent = cap_elem.getparent()
    pos = list(parent).index(cap_elem)
    parent.insert(pos, new_img)

    # 删除旧图片（索引偏移+1）
    if img42_idx >= 0:
        old_p = doc.paragraphs[img42_idx + 1]
        if old_p._element.find('.//' + _qn('w:drawing')) is not None:
            old_p._element.getparent().remove(old_p._element)

    log('  图4-2替换完成')

# ══════════════════════════════════════════════════════════
# 2. 图4-3~4-6：在各节正文结尾、下一节标题前插入流程图
# ══════════════════════════════════════════════════════════
log('--- 插入图4-3~4-6（模块流程图）---')

# 每次插入后段落索引会变，所以用关键词动态查找
flow_configs = [
    # (正文关键词, 下一标题关键词, 图片文件, 图题, 宽度)
    ('家庭组模块', '4.2.2',  'fig4-3_family_flow.png',    '图 4-3  家庭组模块流程图',    8.5),
    ('药品管理模块', '4.2.3', 'fig4-4_medicine_flow.png',  '图 4-4  药品管理模块流程图',  8.5),
    ('帖子模块', '4.2.4',    'fig4-5_post_flow.png',      '图 4-5  帖子串行双审核流程图', 13.0),
    ('系统通知模块', '4.3',   'fig4-6_notify_flow.png',    '图 4-6  系统通知模块流程图',  8.5),
]

for body_key, next_heading_key, img_file, caption, width in flow_configs:
    # 找到下一标题（图片插在它之前）
    next_h_idx = find_idx(next_heading_key)
    if next_h_idx < 0:
        log(f'  未找到下一标题: {next_heading_key}')
        continue

    # 确认该位置前面是模块内容段落
    log(f'  在[{next_h_idx}]"{doc.paragraphs[next_h_idx].text[:30]}"前插入{caption}')

    img_elem = make_img_elem(os.path.join(DIAG, img_file), width)
    cap_elem = make_caption_elem(caption, next_h_idx)

    # 插入图片+图题（在下一标题前）
    anchor = doc.paragraphs[next_h_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)
    parent.insert(pos, cap_elem)
    parent.insert(pos, img_elem)

    log(f'  插入完成: {caption}')

# ══════════════════════════════════════════════════════════
# 3. 图4-7a + 图4-7b ER图：插在"二、实体间主要关系"描述之前
# ══════════════════════════════════════════════════════════
log('--- 插入 ER 图 ---')
er_anchor_idx = find_idx('二、实体间主要关系')
log(f'  ER插入锚点: [{er_anchor_idx}] "{doc.paragraphs[er_anchor_idx].text[:40] if er_anchor_idx>0 else ""}"')

if er_anchor_idx > 0:
    img_a = make_img_elem(os.path.join(DIAG, 'fig4-7a_er_core.png'), 13.5)
    cap_a = make_caption_elem('图 4-7a  E-R 图（核心药品与帖子关系）', er_anchor_idx)
    img_b = make_img_elem(os.path.join(DIAG, 'fig4-7b_er_msg.png'), 13.5)
    cap_b = make_caption_elem('图 4-7b  E-R 图（用户/消息/通知/家庭组申请关系）', er_anchor_idx)

    anchor = doc.paragraphs[er_anchor_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)

    for elem in [img_a, cap_a, img_b, cap_b]:
        parent.insert(pos, elem)
        pos += 1

    log('  ER图插入完成')

# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存！')
