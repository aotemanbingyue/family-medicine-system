# -*- coding: utf-8 -*-
"""
在《初稿 (2).docx》中，在图 4-3、4-4、4-5 的图题「上一行」（用户已准备好格式的段落）直接插入图片。
4-3、4-4 用现有图；4-5 用竖版流程图（已由 gen_all_figs_v4 生成）。
"""
from docx import Document
from docx.shared import Cm
import os, sys

# 微信目录下的文档（若无法写入则保存到桌面）
DOC_PATH = r'd:\SoftWare\微信\文件存储\xwechat_files\wxid_1ya3zyyvfbcf22_970a\msg\file\2026-03\初稿 (2).docx'
OUT_PATH = r'c:\Users\28041\Desktop\初稿 (2)_图43445已插入.docx'  # 备用保存路径
DIAG_DIR = r'd:\SoftWare\Code\bysj\diagrams'

FIGS = [
    ('图 4-3', '家庭组模块流程图', 'fig4-3_family_flow.png', 8.5),
    ('图 4-4', '药品管理模块流程图', 'fig4-4_medicine_flow.png', 8.5),
    ('图 4-5', '帖子串行双审核流程图', 'fig4-5_post_flow.png', 8.5),
]


def log(s):
    sys.stdout.buffer.write((s + '\n').encode('utf-8', 'replace'))


if not os.path.isfile(DOC_PATH):
    log('文档不存在: ' + DOC_PATH)
    sys.exit(1)

doc = Document(DOC_PATH)


def find_caption(key, extra):
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if key in t and extra in t:
            return i
    return -1


for key, extra, fname, width_cm in FIGS:
    cap_idx = find_caption(key, extra)
    if cap_idx < 0:
        log(f'未找到图题：{key} {extra}')
        continue

    if cap_idx < 1:
        log(f'{key} 图题在第一行，无上一行可插入')
        continue

    # 图题上一行 = 用户已准备好格式的段落
    para_above = doc.paragraphs[cap_idx - 1]
    log(f'{key} 图题位置: [{cap_idx}]，上一行: [{cap_idx - 1}]')

    img_path = os.path.join(DIAG_DIR, fname)
    if not os.path.isfile(img_path):
        log(f'  图片不存在：{img_path}')
        continue

    # 清空该段落原有内容，只保留段落格式
    for r in list(para_above.runs):
        r._element.getparent().remove(r._element)

    # 在该段落中插入图片
    run = para_above.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    log(f'  已在图题上一行插入：{fname}')


try:
    doc.save(DOC_PATH)
    log('已保存到原文件。')
except PermissionError:
    doc.save(OUT_PATH)
    log('原路径无法写入，已另存为: ' + OUT_PATH)
