# -*- coding: utf-8 -*-
from docx import Document
import sys

doc = Document(r'd:\SoftWare\Code\bysj\midterm_check.docx')
sys.stdout.buffer.write(f'段落数: {len(doc.paragraphs)}\n'.encode('utf-8','replace'))
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        sys.stdout.buffer.write(f'[{i}] {t}\n'.encode('utf-8','replace'))

# 读表格
sys.stdout.buffer.write(f'\n表格数: {len(doc.tables)}\n'.encode('utf-8','replace'))
for ti, tbl in enumerate(doc.tables):
    sys.stdout.buffer.write(f'\n--- 表格{ti} ---\n'.encode('utf-8','replace'))
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        sys.stdout.buffer.write((' | '.join(cells[:6]) + '\n').encode('utf-8','replace'))
