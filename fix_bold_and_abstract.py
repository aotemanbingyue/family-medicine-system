# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

# ══════════════════════════════════════════════════════════
# 问题1：英文摘要正文前面没有独立的"Abstract："标题
# [53] 是 "ABSTRACT：" 标题行（bold=True），保留
# [54] 是英文摘要正文，检查前面是否已有标题
# ══════════════════════════════════════════════════════════
print('=== 检查摘要区域 ===')
for i in range(46, 58):
    p = doc.paragraphs[i]
    print(f'[{i}] bold={[r.font.bold for r in p.runs]} | {p.text[:60]}')

# ══════════════════════════════════════════════════════════
# 问题2：去掉研究目标4个分点（[85]-[88]）的粗体
# ══════════════════════════════════════════════════════════
for i in range(85, 89):
    para = doc.paragraphs[i]
    for run in para.runs:
        run.font.bold = False
print('去掉研究目标粗体 [85]-[88]')

# ══════════════════════════════════════════════════════════
# 问题3：去掉第二章全部段落的粗体（[114]-[132]）
# 注意：章标题"第二章 相关技术简介"[114] 和节标题"2.1"、"2.2"等 应该保留粗体
# 正文段落去掉粗体
# ══════════════════════════════════════════════════════════
# 章标题、节标题保留粗体，正文段落去掉
heading_markers = ['第二章', '2.1', '2.2', '2.2.1', '2.2.2', '2.2.3', '2.3']

for i in range(114, 133):
    para = doc.paragraphs[i]
    text = para.text.strip()
    is_heading = any(text.startswith(h) for h in heading_markers)
    if not is_heading:
        for run in para.runs:
            run.font.bold = False
        print(f'去粗体 [{i}]: {text[:50]}')
    else:
        print(f'保留粗体（标题）[{i}]: {text[:50]}')

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n修复完成，文件已保存。')
