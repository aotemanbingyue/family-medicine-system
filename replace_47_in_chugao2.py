# -*- coding: utf-8 -*-
"""
在桌面《初稿 (2).docx》中，将图4-7a、图4-7b替换为最新生成的 ER 图。
"""
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, sys

DOC_PATH = r'c:\Users\28041\Desktop\初稿 (2).docx'
DIAG_DIR = r'd:\SoftWare\Code\bysj\diagrams'


def log(s):
    sys.stdout.buffer.write((s + '\n').encode('utf-8', 'replace'))


def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None


doc = Document(DOC_PATH)


def find_cap(key):
    for i, p in enumerate(doc.paragraphs):
        if key in p.text:
            return i
    return -1


TASKS = [
    ('图 4-7a', 'fig4-7a_er_core.png', 13.0),
    ('图 4-7b', 'fig4-7b_er_msg.png', 13.0),
]

for key, fname, wcm in TASKS:
    cap_idx = find_cap(key)
    if cap_idx < 0:
        log(f'未找到图题: {key}')
        continue

    log(f'{key} 图题位置: [{cap_idx}] "{doc.paragraphs[cap_idx].text.strip()[:40]}"')

    # 在图题附近寻找当前图片段落
    img_idx = -1
    for i in range(cap_idx - 6, cap_idx + 7):
        if 0 <= i < len(doc.paragraphs) and has_img(doc.paragraphs[i]):
            img_idx = i
            break
    if img_idx < 0:
        log('  未找到现有图片段落，跳过')
        continue

    # 删除与该图片连续在一起的多余图片（一般不会有，但以防万一）
    # 从 img_idx 向上、向下检查连续的图片段落
    extra = []
    for j in range(img_idx - 1, max(img_idx - 4, -1), -1):
        if has_img(doc.paragraphs[j]):
            extra.append(j)
        else:
            break
    for j in range(img_idx + 1, min(img_idx + 4, len(doc.paragraphs))):
        if has_img(doc.paragraphs[j]):
            extra.append(j)
        else:
            break

    for j in sorted(extra, reverse=True):
        p = doc.paragraphs[j]._element
        p.getparent().remove(p)
        log(f'  删除多余图片段落 [{j}]')

    img_path = os.path.join(DIAG_DIR, fname)
    if not os.path.isfile(img_path):
        log('  图片不存在: ' + img_path)
        continue

    # 用新的图片段落替换原来的 img_idx
    old = doc.paragraphs[img_idx]._element
    parent = old.getparent()
    pos = list(parent).index(old)

    p_new = doc.add_paragraph()
    p_new.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_new.add_run()
    run.add_picture(img_path, width=Cm(wcm))

    e_new = p_new._element
    e_new.getparent().remove(e_new)
    parent.insert(pos, e_new)
    parent.remove(old)

    log(f'  已替换为新图片: {fname} (宽度 {wcm}cm)')


doc.save(DOC_PATH)
log('全部完成，已保存至原文件。')

