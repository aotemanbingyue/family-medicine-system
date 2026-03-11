# -*- coding: utf-8 -*-
"""全文扫描：找出所有图片段落、图题段落、表题段落、表格位置，判断顺序是否正确"""
from docx import Document
from docx.oxml.ns import qn
import sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

def is_table_elem(elem):
    return elem.tag == qn('w:tbl')

# 构建body子元素列表（段落+表格混合）
body = doc.element.body
items = []  # (type, index_in_paragraphs_or_None, text_preview)
para_idx = 0
for child in body:
    if child.tag == qn('w:p'):
        p = doc.paragraphs[para_idx]
        t = p.text.strip()
        img = has_img(p)
        items.append(('img' if img else 'para', para_idx, t[:70]))
        para_idx += 1
    elif child.tag == qn('w:tbl'):
        items.append(('table', None, '[TABLE]'))

sys.stdout.buffer.write(f'总段落: {len(doc.paragraphs)}, 总body子元素: {len(items)}\n'.encode('utf-8','replace'))

# 找出所有图题、表题的位置及其相对图片/表格的关系
problems = []
for i, (typ, pidx, txt) in enumerate(items):
    is_fig_cap = typ == 'para' and pidx is not None and ('图 ' in txt or '图5-' in txt or '图4-' in txt or '图3-' in txt) and any(c.isdigit() for c in txt[:10])
    is_tbl_cap = typ == 'para' and pidx is not None and ('表 ' in txt or '表6-' in txt or '表4-' in txt) and any(c.isdigit() for c in txt[:10])
    
    if is_fig_cap:
        # 检查前一个非空item是否是图片
        prev_img = False
        for j in range(i-1, max(i-4,-1), -1):
            pt, pp, ptxt = items[j]
            if pt == 'img': prev_img = True; break
            if pt == 'para' and ptxt: break
        if not prev_img:
            problems.append(('FIG_CAP_ABOVE', pidx, txt))
            sys.stdout.buffer.write(f'  [图题在上] [{pidx}] {txt}\n'.encode('utf-8','replace'))
        else:
            sys.stdout.buffer.write(f'  [图题OK]   [{pidx}] {txt}\n'.encode('utf-8','replace'))
    
    if is_tbl_cap:
        # 检查前一个非空item是否是表格
        prev_tbl = False
        for j in range(i-1, max(i-4,-1), -1):
            pt, pp, ptxt = items[j]
            if pt == 'table': prev_tbl = True; break
            if pt == 'para' and ptxt: break
        if not prev_tbl:
            problems.append(('TBL_CAP_ABOVE', pidx, txt))
            sys.stdout.buffer.write(f'  [表题在上] [{pidx}] {txt}\n'.encode('utf-8','replace'))
        else:
            sys.stdout.buffer.write(f'  [表题OK]   [{pidx}] {txt}\n'.encode('utf-8','replace'))

sys.stdout.buffer.write(f'\n共发现 {len(problems)} 个位置错误\n'.encode('utf-8','replace'))
