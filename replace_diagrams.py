# -*- coding: utf-8 -*-
"""
替换 Word 中已有图片（通过删除旧图片段落+图题段落，重新插入），
并在 4.3.1 插入两张 ER 图（fig4-7a, fig4-7b）
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, os, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
DIAG = r'd:\SoftWare\Code\bysj\diagrams'

def log(s):
    sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

# ─────────────────────────────────────────────────────────
# 工具函数
# ─────────────────────────────────────────────────────────
def find_idx(*keys, start=0):
    for i in range(start, len(doc.paragraphs)):
        try:
            if all(k in doc.paragraphs[i].text for k in keys):
                return i
        except: pass
    return -1

def para_has_image(para):
    return para._element.find('.//' + qn('w:drawing')) is not None or \
           para._element.find('.//' + qn('v:shape')) is not None

def remove_para(idx):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

def make_centered_image_para(img_path, width_cm):
    """创建居中的图片段落元素"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def make_caption_para(text, ref_para):
    """创建居中加粗图题段落元素"""
    new_p = OxmlElement('w:p')
    # 复制段落格式
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        new_p.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b'); rPr.append(b)
    # 字号 10.5pt = 小四
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '21'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '21'); rPr.append(szCs)
    new_r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(t)
    new_p.append(new_r)
    return new_p

def insert_after(anchor_idx, *elems):
    """在 anchor_idx 段落之后依次插入若干元素"""
    anchor = doc.paragraphs[anchor_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor) + 1
    for elem in elems:
        parent.insert(pos, elem)
        pos += 1

def insert_before(anchor_idx, *elems):
    anchor = doc.paragraphs[anchor_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)
    for elem in elems:
        parent.insert(pos, elem)
        pos += 1

# ─────────────────────────────────────────────────────────
# 查找并删除旧图片段落（含图题）
# ─────────────────────────────────────────────────────────
def replace_image_block(caption_keyword, new_img_path, width_cm, search_range=5):
    """
    找到包含 caption_keyword 的图题段落，
    向前查找图片段落并删除，再插入新图片+图题。
    """
    cap_idx = find_idx(caption_keyword)
    if cap_idx < 0:
        log(f'  未找到图题: {caption_keyword}')
        return

    # 向前找图片段落（最多 search_range 行）
    img_idx = -1
    for i in range(cap_idx - 1, max(cap_idx - search_range - 1, -1), -1):
        if para_has_image(doc.paragraphs[i]):
            img_idx = i
            break

    ref_para = doc.paragraphs[cap_idx]

    # 先获取旧图题文字，作为新图题
    cap_text = doc.paragraphs[cap_idx].text.strip()

    # 创建新元素
    new_img_elem  = make_centered_image_para(new_img_path, width_cm)
    new_cap_elem  = make_caption_para(cap_text, ref_para)

    # 插入在图题之前（新图片 + 新图题）
    cap_anchor = doc.paragraphs[cap_idx]._element
    parent = cap_anchor.getparent()
    cap_pos = list(parent).index(cap_anchor)

    parent.insert(cap_pos, new_cap_elem)
    parent.insert(cap_pos, new_img_elem)

    # 删除旧图题（现在 cap_idx+2）和旧图片
    # 先重新找旧图题（文字相同），避免索引偏移
    old_cap_idx = find_idx(caption_keyword, start=cap_idx + 2)
    if old_cap_idx > 0:
        remove_para(old_cap_idx)
    if img_idx >= 0:
        # 重新找旧图片段落
        for i in range(cap_idx, cap_idx + 5):
            if para_has_image(doc.paragraphs[i]):
                remove_para(i)
                break

    log(f'  替换完成: {cap_text}')


# ══════════════════════════════════════════════════════════
# 1. 替换 图4-2 ~ 图4-6
# ══════════════════════════════════════════════════════════
log('--- 替换图4-2~4-6 ---')
replace_image_block('图 4-2', os.path.join(DIAG, 'fig4-2_module_tree.png'), 14.0)
replace_image_block('图 4-3', os.path.join(DIAG, 'fig4-3_family_flow.png'), 8.5)
replace_image_block('图 4-4', os.path.join(DIAG, 'fig4-4_medicine_flow.png'), 8.5)
replace_image_block('图 4-5', os.path.join(DIAG, 'fig4-5_post_flow.png'), 14.0)
replace_image_block('图 4-6', os.path.join(DIAG, 'fig4-6_notify_flow.png'), 8.5)

# ══════════════════════════════════════════════════════════
# 2. 替换旧 ER 图 (图4-7) → 插入图4-7a + 图4-7b
# ══════════════════════════════════════════════════════════
log('--- 替换 ER 图 ---')
er_cap_idx = find_idx('图 4-7')
log(f'  旧ER图题位置: [{er_cap_idx}]')

if er_cap_idx > 0:
    ref_para = doc.paragraphs[er_cap_idx]
    parent = ref_para._element.getparent()
    pos = list(parent).index(ref_para._element)

    # 构造两张新图的元素
    img_a = make_centered_image_para(os.path.join(DIAG, 'fig4-7a_er_core.png'), 13.5)
    cap_a = make_caption_para('图 4-7a  E-R 图（核心药品与帖子关系）', ref_para)
    img_b = make_centered_image_para(os.path.join(DIAG, 'fig4-7b_er_msg.png'), 13.5)
    cap_b = make_caption_para('图 4-7b  E-R 图（用户/消息/通知/家庭组申请关系）', ref_para)

    # 在旧图题处依次插入
    for elem in [img_a, cap_a, img_b, cap_b]:
        parent.insert(pos, elem)
        pos += 1

    # 删除旧图题
    remove_para(er_cap_idx + 4)

    # 删除旧图片
    for i in range(er_cap_idx + 3, er_cap_idx + 8):
        if i < len(doc.paragraphs) and para_has_image(doc.paragraphs[i]):
            remove_para(i)
            break

    log('  ER 图替换完成')

# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存。')
