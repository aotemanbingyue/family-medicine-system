# -*- coding: utf-8 -*-
"""
1. 给第五章所有图片后插入居中不加粗图题
2. 修复末尾：总结加章标题、删参考文献占位行、清理多余空行
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

def find_idx(*keys, start=0):
    for i in range(start, len(doc.paragraphs)):
        if all(k in doc.paragraphs[i].text for k in keys):
            return i
    return -1

def set_text(idx, text, bold=False):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''; r.font.bold = False
    if p.runs:
        p.runs[0].text = text; p.runs[0].font.bold = bold
    else:
        r = p.add_run(text); r.font.bold = bold

def remove_para(idx):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

def make_caption_elem(text, ref_para):
    """居中、不加粗的图题段落元素"""
    new_p = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); new_p.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(qn('w:val'), 'center')

    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # 不加粗，字号10.5pt（小四）
    sz = OxmlElement('w:sz');   sz.set(qn('w:val'),  '21'); rPr.append(sz)
    szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), '21'); rPr.append(szCs)
    new_r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(t); new_p.append(new_r)
    return new_p

def insert_after_elem(anchor_elem, new_elem):
    parent = anchor_elem.getparent()
    pos = list(parent).index(anchor_elem)
    parent.insert(pos + 1, new_elem)

# ══════════════════════════════════════════════════════════
# 1. 给第五章图片加图题
#    图片分布：
#    5.1: [245][247][249][251]  → 图5-1 ~ 图5-4
#    5.2: [259][261][263]       → 图5-5 ~ 图5-7
#    5.3: [270][272][274][276]  → 图5-8 ~ 图5-11
#    5.4: [284][286][288]       → 图5-12 ~ 图5-14
#
#    图题内容按模块功能顺序命名
# ══════════════════════════════════════════════════════════
log('=== 给第五章图片加图题 ===')

# 图题映射：(节标题关键词, [(图序号, 图题文字), ...])
captions_by_section = [
    ('5.1.2', [
        ('图 5-1', '家庭组协作页面'),
        ('图 5-2', '家庭药箱管理页面'),
        ('图 5-3', '申请加入家庭组页面'),
        ('图 5-4', '用户管理员审核申请页面'),
    ]),
    ('5.2.2', [
        ('图 5-5', '全局标准药品库检索页面'),
        ('图 5-6', '家庭药箱药品列表页面'),
        ('图 5-7', '药品管理员审核药品申请页面'),
    ]),
    ('5.3.2', [
        ('图 5-8',  '共享广场帖子列表页面'),
        ('图 5-9',  '帖子发布页面'),
        ('图 5-10', '帖子管理员审核页面'),
        ('图 5-11', '用户采纳药品入库页面'),
    ]),
    ('5.4.2', [
        ('图 5-12', '用户消息中心页面'),
        ('图 5-13', '系统公告与医学小贴士推送页面'),
        ('图 5-14', '私信沟通页面'),
    ]),
]

# 检查是否已有图题（避免重复插入）
def has_caption_after(img_para_elem):
    parent = img_para_elem.getparent()
    idx = list(parent).index(img_para_elem)
    # 看下一个非空段落是否包含"图 5-"
    for j in range(idx+1, min(idx+3, len(parent))):
        child = parent[j]
        if child.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in child.findall('.//' + qn('w:t')))
            if '图 5-' in text:
                return True
            if text.strip():  # 遇到非空非图题段落，停止
                break
    return False

for sec_key, cap_list in captions_by_section:
    sec_idx = find_idx(sec_key)
    if sec_idx < 0:
        log(f'  未找到节: {sec_key}'); continue

    # 收集该节内的图片段落
    img_paras = []
    for i in range(sec_idx+1, min(sec_idx+25, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        t = p.text.strip()
        # 遇到下一节标题停止
        if t and not has_img(p) and any(
            t.startswith(x) for x in ['5.2','5.3','5.4','5.5','第六章','6.']):
            break
        if has_img(p):
            img_paras.append(i)

    log(f'  {sec_key}: 找到{len(img_paras)}张图片, 图题数{len(cap_list)}')

    for k, img_idx in enumerate(img_paras):
        if k >= len(cap_list): break
        num, title = cap_list[k]
        img_elem = doc.paragraphs[img_idx]._element

        if has_caption_after(img_elem):
            log(f'    [{img_idx}] 已有图题，跳过')
            continue

        cap_text = f'{num}  {title}'
        cap_elem = make_caption_elem(cap_text, doc.paragraphs[img_idx])
        insert_after_elem(img_elem, cap_elem)
        log(f'    [{img_idx}] 插入图题: {cap_text}')

# ══════════════════════════════════════════════════════════
# 2. 修复末尾内容
# ══════════════════════════════════════════════════════════
log('\n=== 修复末尾 ===')

# 2a. "总结与展望"加章标题前缀
idx_summary = find_idx('总结与展望')
log(f'"总结与展望"位置: [{idx_summary}]')
if idx_summary >= 0:
    cur = doc.paragraphs[idx_summary].text.strip()
    if not cur.startswith('第'):
        set_text(idx_summary, '第七章  总结与展望', bold=True)
        log('  已加"第七章"前缀')

# 2b. 删除参考文献占位行
idx_ref_hint = find_idx('内容：参考的书籍')
log(f'参考文献占位行: [{idx_ref_hint}]')
if idx_ref_hint >= 0:
    remove_para(idx_ref_hint)
    log('  已删除参考文献占位行')

# 2c. 清理总结与展望中多余空行（连续3个以上空行压缩为1个）
idx_summary2 = find_idx('第七章  总结与展望')
if idx_summary2 < 0: idx_summary2 = find_idx('总结与展望')
idx_ref = find_idx('参考文献')
log(f'总结范围: [{idx_summary2}] ~ [{idx_ref}]')
if idx_summary2 >= 0 and idx_ref > idx_summary2:
    i = idx_summary2 + 1
    consecutive_empty = 0
    while i < find_idx('参考文献'):
        t = doc.paragraphs[i].text.strip()
        if not t:
            consecutive_empty += 1
            if consecutive_empty > 1:
                remove_para(i)
                continue
        else:
            consecutive_empty = 0
        i += 1

# 2d. 第六章测试部分清理多余空行（6.2节内大量空行）
idx_62 = find_idx('6.2 系统功能测试')
idx_63 = find_idx('6.3 系统性能测试')
log(f'6.2~6.3范围: [{idx_62}] ~ [{idx_63}]')
if idx_62 >= 0 and idx_63 > idx_62:
    i = idx_62 + 1
    consecutive_empty = 0
    while i < find_idx('6.3 系统性能测试'):
        t = doc.paragraphs[i].text.strip()
        if not t:
            consecutive_empty += 1
            if consecutive_empty > 1:
                remove_para(i)
                continue
        else:
            consecutive_empty = 0
        i += 1
    log('  6.2节多余空行已清理')

# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存！')
