# -*- coding: utf-8 -*-
"""
1. 替换 4.3 核心模块设计内容
2. 补充 4.4.1 ER图完整描述
3. 将所有数据库表格改为三线表
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

# ── 工具函数 ──────────────────────────────────────────────

def make_para(text, ref_para, bold=False):
    new_p = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    new_r = OxmlElement('w:r')
    ref_runs = ref_para._element.findall(qn('w:r'))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))
    rPr = new_r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        new_r.insert(0, rPr)
    b = rPr.find(qn('w:b'))
    if b is not None:
        rPr.remove(b)
    if bold:
        rPr.append(OxmlElement('w:b'))
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p

def find_idx(*keys):
    for i, p in enumerate(doc.paragraphs):
        try:
            if all(k in p.text for k in keys):
                return i
        except: pass
    return -1

def replace_range(start_idx, end_idx, items_bold):
    """items_bold: list of (text, bold)"""
    ref = doc.paragraphs[start_idx]
    elems = [doc.paragraphs[i]._element for i in range(start_idx, end_idx + 1)]
    parent = elems[0].getparent()
    insert_pos = list(parent).index(elems[0])
    for i, (text, bold) in enumerate(items_bold):
        parent.insert(insert_pos + i, make_para(text, ref, bold))
    for e in elems:
        parent.remove(e)

def insert_before_idx(idx, items_bold, ref_idx=None):
    ref = doc.paragraphs[ref_idx if ref_idx else idx]
    anchor = doc.paragraphs[idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)
    for i, (text, bold) in enumerate(items_bold):
        parent.insert(pos + i, make_para(text, ref, bold))

# ══════════════════════════════════════════════════════════
# 1. 替换 4.3 核心模块设计
#    [189] = 4.3.1 用户加入家庭组模块设计
#    [190] = 模块功能总体描述...（占位提示）
#    [191] = 4.3.2 帖子模块设计
#    [192] = 4.3.4 系统通知模块设计
# ══════════════════════════════════════════════════════════
print('--- 替换 4.3 ---')
idx_431 = find_idx('4.3.1', '用户加入家庭组')
idx_432 = find_idx('4.3.4', '系统通知')
print(f'4.3.1~4.3.4: [{idx_431}]-[{idx_432}]')

new_43 = [
    ('4.3.1  家庭组模块设计', True),
    ('家庭组模块是系统的基础支撑模块，负责用户身份管理与家庭组协作关系的建立。其核心流程如下：用户访问系统后首先完成注册，填写用户名与密码后账号激活；注册成功后登录，系统根据 role 字段分配角色权限；普通用户登录后可创建家庭组（系统生成唯一 family_id）或提交加入申请；加入申请提交后进入"待审核"状态，由用户管理员在后台审批，审核通过后系统自动更新用户的家庭归属；家庭组成员可查看同组成员的药箱信息，实现家庭内协作管理。用户管理员还可在后台查看、启用/禁用用户账号及重置密码，保障平台用户质量。', False),
    ('4.3.2  药品管理模块设计', True),
    ('药品管理模块是系统的核心业务模块，负责全局标准药库的维护和家庭药箱的全生命周期管理。其核心流程如下：药品管理员对全局标准药库进行分类维护（增删改）并支持一键同步更新；普通用户可在标准药库中检索目标药品，点击"添加至家庭药箱"后，系统创建一条 audit_status 为"待审核"的家庭药箱记录；药品管理员在审核列表中对该申请进行审核，审核通过后记录状态变更为"已通过"，用户即可正式管理该药品；系统自动检测每条药品记录的有效期，对距有效期不足 30 天的记录在用户首页以红点角标提示，并在消息中心发送过期提醒；用户可为每条药品设置每日服药提醒时间与剂量备注，到时系统在服药提醒页面进行展示。', False),
    ('4.3.3  帖子模块设计', True),
    ('帖子模块实现用户间的药品信息交流与闲置药品流转撮合，采用串行双审核机制保障帖子内容与药品信息的双重合规性。其核心流程如下：用户填写帖子信息并绑定一条家庭药箱记录（作为药品凭证）后提交发布，帖子初始状态为"待帖子审核"；帖子管理员对帖子内容进行合规性审核，审核通过后帖子进入"待药品审核"状态，并自动通知药品管理员；药品管理员对帖子引用的药品信息进行专业审核，两级均通过后帖子正式在共享广场公开；其他用户可浏览已审核帖子，点击"加入我的家庭药箱"后，系统以该帖子药品信息为模板为用户创建新的家庭药箱记录（再次进入药品审核流程）；审核驳回时，系统将驳回理由推送至发帖用户的消息中心。', False),
    ('4.3.4  系统通知模块设计', True),
    ('系统通知模块负责将平台重要信息主动推送至用户，并提供用户与管理员之间的私信沟通渠道。其核心流程如下：系统管理员在后台撰写并发布系统公告或每日医学小贴士，发布后所有用户均可在公告列表与首页卡片中查看；药品到期提醒由系统自动生成，定期扫描家庭药箱中距有效期不足 30 天的记录，并在服药提醒页面汇总展示；用户可通过"联系管理员"入口向任意管理员发送私信，管理员在私聊收件箱中集中查看与回复；消息中心（用户信箱）统一汇聚私信、公告、小贴士、帖子审核反馈等多类消息，以未读角标提示用户及时查阅。', False),
]
replace_range(idx_431, idx_432, new_43)
print('4.3替换完成')

# ══════════════════════════════════════════════════════════
# 2. 补充 4.4.1 ER图描述
#    在现有3条关系描述前，插入完整的实体+关系说明
# ══════════════════════════════════════════════════════════
print('--- 补充 4.4.1 ER图 ---')
idx_441_first = find_idx('一个用户（User）只能归属于一个家庭组')
print(f'4.4.1首行: [{idx_441_first}]')

er_intro = [
    ('本系统的 E-R 图包含 9 个核心实体，各实体的属性及实体间的关联关系说明如下。', False),
    ('一、核心实体说明', False),
    ('用户（User）：系统的核心参与者，主要属性包括用户名、邮箱、密码、角色（role）、家庭组标识（family_id）和账号状态（is_active）。角色字段区分普通用户与四类管理员，家庭组标识标记该用户所属家庭。', False),
    ('全局标准药库（GlobalMedicine）：由药品管理员维护的标准化药品字典，主要属性包括药品名称、分类、生产厂家、功效说明及软删除标记。', False),
    ('家庭药箱（FamilyMedicine）：记录用户个人药品库存，主要属性包括药品名称、数量、有效期、审核状态（audit_status）、服药提醒开关与时间，以及软删除标记。', False),
    ('共享帖子（SharePost）：用户发布的药品转让信息，主要属性包括标题、正文、帖子审核状态（status）、药品二级审核状态（medicine_audit_status）及软删除标记。', False),
    ('家庭组加入申请（FamilyJoinRequest）：记录用户申请加入家庭组的过程，主要属性包括目标家庭组标识、申请状态（pending/approved/rejected）、审核备注及审核时间。', False),
    ('私信消息（PrivateMessage）：记录用户与管理员之间的私信内容，主要属性包括发送方、接收方、消息正文、发送时间和已读时间。', False),
    ('系统公告（SystemAnnouncement）：由系统管理员发布的全站公告，主要属性包括标题、正文和发布时间。', False),
    ('医学小贴士（MedicalTip）：由系统管理员每日发布的健康知识，主要属性包括发布日期（每日唯一）、标题和正文。', False),
    ('帖子药品采纳记录（SharePostMedicineAdoption）：记录用户从帖子中采纳药品并加入家庭药箱的操作，作为关联实体连接用户、帖子与家庭药箱。', False),
    ('二、实体间主要关系', False),
]
insert_before_idx(idx_441_first, er_intro, ref_idx=idx_441_first)
print('ER图描述插入完成')

# 对原有3条关系描述（现在索引已变，重新找）去掉代码类名称
for i, p in enumerate(doc.paragraphs):
    try:
        t = p.text.strip()
        if 'family_id 标识' in t and '1:N' in t:
            for run in p.runs: run.font.bold = False
        if 'FamilyMedicine' in t and 'GlobalMedicine' in t and '数据校准' in t:
            for run in p.runs: run.font.bold = False
        if 'SharePost' in t and 'FamilyMedicine' in t and '药品真实性' in t:
            for run in p.runs: run.font.bold = False
    except: pass

# ══════════════════════════════════════════════════════════
# 3. 将所有数据库表格改为三线表
#    三线表：顶线（粗）+ 表头底线（细）+ 底线（粗）
#    去掉内部所有纵向格线和中间横线
# ══════════════════════════════════════════════════════════
print('--- 改三线表 ---')

def make_border(val='single', sz='12', color='000000'):
    b = OxmlElement('w:top')  # placeholder, caller sets tag
    b.set(qn('w:val'), val)
    b.set(qn('w:sz'), sz)
    b.set(qn('w:color'), color)
    b.set(qn('w:space'), '0')
    return b

def set_three_line_table(tbl):
    """将一个 tbl 元素改为三线表"""
    rows = tbl.findall(qn('w:tr'))
    if not rows:
        return
    n_rows = len(rows)

    # 先去掉整表的格线（tblPr 里的 tblBorders）
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # 删除已有 tblBorders
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    # 设置整表无边框
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)

    for r_idx, tr in enumerate(rows):
        cells = tr.findall(qn('w:tc'))
        is_first = (r_idx == 0)
        is_header = (r_idx == 0)     # 第一行 = 表头
        is_second = (r_idx == 1)     # 第二行 = 表头下方第一行数据
        is_last = (r_idx == n_rows - 1)

        for tc in cells:
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            # 删除旧 tcBorders 和 shd
            old_b = tcPr.find(qn('w:tcBorders'))
            if old_b is not None:
                tcPr.remove(old_b)
            old_shd = tcPr.find(qn('w:shd'))
            if old_shd is not None:
                tcPr.remove(old_shd)

            tcBorders = OxmlElement('w:tcBorders')

            def add_side(tag, val, sz, color='000000'):
                s = OxmlElement(f'w:{tag}')
                s.set(qn('w:val'), val)
                s.set(qn('w:sz'), sz)
                s.set(qn('w:color'), color)
                s.set(qn('w:space'), '0')
                tcBorders.append(s)

            # 顶线：第一行顶部 粗线
            if is_first:
                add_side('top', 'single', '18')      # 1.5pt 粗线
            else:
                add_side('top', 'none', '0')

            # 表头底线：第一行底部（= 第二行顶部）细线
            if is_header:
                add_side('bottom', 'single', '6')    # 0.5pt 细线
            elif is_last:
                add_side('bottom', 'single', '18')   # 底线粗
            else:
                add_side('bottom', 'none', '0')

            # 无纵向格线
            add_side('left', 'none', '0')
            add_side('right', 'none', '0')
            add_side('insideH', 'none', '0')
            add_side('insideV', 'none', '0')

            tcPr.append(tcBorders)

            # 去掉表头灰色背景（改为白色）
            if is_header:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFFFF')
                tcPr.append(shd)

# 找到文档中所有表格，对名称中含"app_"的表格（数据库表）进行三线化
# 找到 4.4.2 节之后的所有 tbl 元素
body = doc.element.body
tables_converted = 0
for elem in body:
    if elem.tag == qn('w:tbl'):
        set_three_line_table(elem)
        tables_converted += 1
print(f'共将 {tables_converted} 张表格改为三线表')

# ══════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n全部完成，文件已保存。')
