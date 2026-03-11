# -*- coding: utf-8 -*-
from docx import Document

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

# ── 1. 删除目录页的"内容：自动生成的目录"提示行 [58]
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if '内容：' in t and '自动' in t and '目录' in t:
        para._element.getparent().remove(para._element)
        print(f'删除目录提示行 [{i}]: {t[:40]}')
        break

# ── 2. 去掉论文结构安排各行里多余的粗体run
#    论文结构安排段落格式：(序号) 粗体章名：正文说明
#    正确格式应该是：章号+章名 粗体，冒号后正文不粗体
#    但目前 [102]-[105] 部分run是True
#    把论文结构安排所有行统一处理：整段去掉粗体（这些是正文描述行，不是标题）
structure_keywords = ['第一章', '第二章', '第三章', '第四章', '第五章', '第六章', '第七章',
                      '绪论', '相关技术', '需求分析', '系统设计', '系统实现', '系统测试', '总结']

for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if any(k in t for k in ['第一章', '第二章', '第三章', '第四章', '第五章', '第六章', '第七章']):
        if any(k in t for k in ['绪论', '相关技术', '需求', '设计', '实现', '测试', '总结', '展望']):
            for run in para.runs:
                run.font.bold = False
            print(f'去粗体（论文结构行）[{i}]: {t[:60]}')

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n修复完成，文件已保存。')

# 最终验证
print('\n=== 最终验证第一章 ===')
doc2 = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    if i >= 57 and i <= 110:
        bold = [r.font.bold for r in para.runs]
        print(f'[{i}] bold={bold[:2]} | {text[:70]}')
