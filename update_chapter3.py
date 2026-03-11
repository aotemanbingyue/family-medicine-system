# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def make_para(text, ref_para):
    new_p = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    new_r = OxmlElement('w:r')
    ref_runs = ref_para._element.findall(qn('w:r'))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))
    # 强制去掉粗体
    rPr = new_r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        new_r.insert(0, rPr)
    b = rPr.find(qn('w:b'))
    if b is not None:
        rPr.remove(b)
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p

def replace_range(start_idx, end_idx, texts):
    ref = doc.paragraphs[start_idx]
    elems = [doc.paragraphs[i]._element for i in range(start_idx, end_idx + 1)]
    parent = elems[0].getparent()
    insert_pos = list(parent).index(elems[0])
    for i, text in enumerate(texts):
        parent.insert(insert_pos + i, make_para(text, ref))
    for e in elems:
        parent.remove(e)

def remove_by_content(*keywords):
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if all(kw in t for kw in keywords):
            para._element.getparent().remove(para._element)
            print(f'  删除 [{i}]: {t[:40]}')
            return True
    return False

def set_para_not_bold(idx):
    para = doc.paragraphs[idx]
    for run in para.runs:
        run.font.bold = False

# ══════════════════════════════════════════════════════════
# 3.1 业务流程分析
# 删除"内容："提示行，替换 [157]-[161]
# ══════════════════════════════════════════════════════════
print('--- 3.1 业务流程 ---')
remove_by_content('内容：', '业务流程')
# 重新定位，[157]-[161] 是业务流程内容（5段）
# 替换为4段干净版本
replace_range(157, 161, [
    '在传统的家庭药品管理模式中，成员之间信息不对称，管理盲区明显。本系统通过多个核心业务流程，将用户操作、管理员审核与系统状态有机串联，形成"录入—审核—流转"的闭环业务体系。',
    '药品准入流程：普通用户在添加或编辑家庭药箱中的药品时，系统默认将该条目标记为"待审核"状态；药品管理员在后台对录入信息进行规范性审核，确认信息的标准化与准确性后，用户方可正式使用该药品记录。',
    '共享广场转让流程：用户发布的药品转让帖需先经帖子管理员审核内容合规性，审核通过后进入药品管理员的二次审核（审核帖子内引用药品的真实性），两级审核均通过后帖子方可公开展示，其他用户可进一步将帖子内药品添加至自己的家庭药箱。',
    '家庭协作审批流程：为防止陌生用户随意加入特定家庭组，用户提交加入申请后须经用户管理员审核，审核通过后系统才会更新该用户的家庭组归属。',
])
print('3.1替换完成')

# ══════════════════════════════════════════════════════════
# 3.2 功能需求分析
# 删除"内容："提示行 [166]
# 替换 3.2.1 [173]-[176]（4段）
# 替换 3.2.2 [177-开头介绍] 和 [178]-[182]（管理员4条）
# ══════════════════════════════════════════════════════════
print('--- 3.2 功能需求 ---')
remove_by_content('内容：', '用例')
print('删除3.2内容提示行完成')

# 3.2.1 现在索引已变，重新找
for i, p in enumerate(doc.paragraphs):
    if '用户是系统的主要服务对象' in p.text:
        idx_321_intro = i
        break
for i, p in enumerate(doc.paragraphs):
    if '广场交互：支持发布' in p.text:
        idx_321_end = i
        break
print(f'3.2.1范围 [{idx_321_intro}]-[{idx_321_end}]')
replace_range(idx_321_intro, idx_321_end, [
    '普通用户是系统的主要使用群体，其核心功能需求包括以下几类：',
    '药箱管理功能：支持药品信息的增删改查，系统自动识别距有效期不足 30 天的药品，在首页以角标提醒方式显示，并支持设置每日服药提醒时间与剂量备注。',
    '共享广场功能：支持发布药品转让帖子并附带药品引用；支持浏览已审核帖子，并可一键将帖子内药品添加至自己的家庭药箱；支持通过私信功能联系管理员。',
    '家庭组功能：支持创建或申请加入家庭组，加入申请须经用户管理员审批；支持查看同家庭组成员的药箱信息，促进家庭内药品协作管理。',
])
print('3.2.1替换完成')

# 3.2.2 重新找
for i, p in enumerate(doc.paragraphs):
    if '为保证平台的秩序' in p.text:
        idx_322_intro = i
        break
for i, p in enumerate(doc.paragraphs):
    if '系统管理员：拥有最高权限' in p.text:
        idx_322_end = i
        break
print(f'3.2.2范围 [{idx_322_intro}]-[{idx_322_end}]')
replace_range(idx_322_intro, idx_322_end, [
    '为保证平台内容与数据的安全合规，系统设有四类管理员角色，各司其职：',
    '药品管理员：负责维护全局标准药库（增删改查与一键同步），以及对用户提交的家庭药品录入申请和共享帖子内引用药品进行专业审核。',
    '帖子管理员：负责对共享广场中所有待发布帖子进行内容合规审核，对违规内容进行拦截、驳回或处理，并将审核结果反馈至用户消息中心。',
    '用户管理员：负责处理家庭组加入申请，并对用户账号进行日常管理（查看信息、启用/禁用账号、重置密码）。',
    '系统管理员：拥有最高权限，负责发布系统公告、每日医学小贴士，以及对系统全局数据进行统筹维护。',
])
print('3.2.2替换完成')

# ══════════════════════════════════════════════════════════
# 3.3 非功能性需求
# 删除"内容："提示行，替换内容并去掉粗体
# ══════════════════════════════════════════════════════════
print('--- 3.3 非功能性需求 ---')
remove_by_content('内容：', '非功能')
# 重新找
for i, p in enumerate(doc.paragraphs):
    if '稳定性：由于涉及' in p.text:
        idx_33_start = i
        break
for i, p in enumerate(doc.paragraphs):
    if '可扩展性：系统代码' in p.text:
        idx_33_end = i
        break
print(f'3.3范围 [{idx_33_start}]-[{idx_33_end}]')
replace_range(idx_33_start, idx_33_end, [
    '稳定性：由于系统涉及家庭用药安全，需保证服务的持续可用性。Django 框架的成熟性与 MySQL 数据库的事务保障，能够为日常多用户并发访问提供稳定支撑。',
    '易用性：考虑到用户群体包含老年人等非技术用户，前端界面设计需简洁直观，关键操作（如审核结果、药品到期提醒）应以明确的文字和颜色区分反馈给用户。',
    '可扩展性：系统采用模块化设计，各功能模块通过 Django 应用结构隔离，权限控制以装饰器方式实现，便于未来新增角色或扩展功能模块而无需大幅重构。',
    '安全性：系统借助 Django 内置的 CSRF 防护、密码哈希存储及自定义角色校验装饰器，形成多层次安全防线，有效防止常见 Web 攻击与越权访问。',
])
print('3.3替换完成')

# ══════════════════════════════════════════════════════════
# 3.4 可行性分析
# 删除"内容："提示行
# ══════════════════════════════════════════════════════════
print('--- 3.4 可行性分析 ---')
remove_by_content('内容：', '操作')
print('3.4内容提示行处理完成')

# ══════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n第三章全部更新完成，文件已保存。')
