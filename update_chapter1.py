# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def make_para(text, ref_para):
    new_p = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    new_r = OxmlElement('w:r')
    ref_runs = ref_para._element.findall(qn('w:r'))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p

def replace_range(start_idx, end_idx, new_texts):
    ref = doc.paragraphs[start_idx]
    elems = [doc.paragraphs[i]._element for i in range(start_idx, end_idx + 1)]
    parent = elems[0].getparent()
    insert_pos = list(parent).index(elems[0])
    for i, text in enumerate(new_texts):
        parent.insert(insert_pos + i, make_para(text, ref))
    for e in elems:
        parent.remove(e)

def replace_single(idx, new_text):
    replace_range(idx, idx, [new_text])

# ══════════════════════════════════════════════════════════
# 第一步：替换研究背景 [64]-[68]
# ══════════════════════════════════════════════════════════
bg = [
    '随着经济的持续发展和人民健康意识的不断增强，越来越多的家庭开始自行储备常用药品，家庭药箱已成为日常健康管理的重要组成部分。然而，在家庭药品管理规模扩大的同时，也暴露出一系列隐患与管理盲区，亟需加以重视和改善[7]。',
    '当前家庭药品管理主要面临以下三类问题：一是安全隐患，药品使用不当、存储条件不达标、过期药品未能及时筛查；二是信息缺失，家庭内部缺乏"统一台账"的管理机制，无法系统追踪库存和到期情况；三是资源浪费，闲置药品缺乏有效流转渠道，造成重复购买和药品积压。',
    '在移动互联网与信息化技术广泛普及的背景下，利用软件平台提升家庭药品管理的信息化水平、为用户提供规范化用药指导，并构建面向家庭间的药品共享机制，已具备充分的技术条件和现实需求。为此，本文设计并实现了一套面向家庭用户的医药管理与共享平台，旨在提升家庭药品管理的规范性与安全性，并提供跨家庭的闲置药品互助渠道。',
]
replace_range(64, 68, bg)
print('步骤1完成：研究背景替换')

# ══════════════════════════════════════════════════════════
# 第二步：删除研究背景节内的"内容："提示行 [62]（现在索引已变，重新找）
# ══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if '内容：根据' in para.text and '研究' in para.text:
        para._element.getparent().remove(para._element)
        print(f'删除提示行 [{i}]')
        break

# ══════════════════════════════════════════════════════════
# 第三步：替换研究现状 - 找到"内容："提示行并删除
# ══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if '内容：' in para.text and '国内外' in para.text:
        para._element.getparent().remove(para._element)
        print(f'删除研究现状提示行 [{i}]')
        break

# ══════════════════════════════════════════════════════════
# 第四步：替换研究目标 - 找到"内容："提示行并删除，替换目标段落
# ══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if '内容：' in para.text and '研究目标' in para.text:
        para._element.getparent().remove(para._element)
        print(f'删除研究目标提示行 [{i}]')
        break

# 找到研究目标的实际内容段（含"角色权限"或"本研究主要围绕"），替换为规范内容
target_start = -1
target_end = -1
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if '本研究主要围绕' in t or '角色权限架构设计' in t:
        target_start = i
    if target_start > 0 and ('软删除' in t or '交互优化' in t or 'context_processors' in t):
        target_end = i
        break

print(f'研究目标范围：[{target_start}] - [{target_end}]')

if target_start > 0 and target_end > 0:
    goal_texts = [
        '本研究主要围绕系统的架构设计、功能实现及安全性与可扩展性展开，具体研究内容如下：',
        '（1）角色权限体系设计。基于 Django 自定义用户模型，设计普通用户、用户管理员、药品管理员、帖子管理员和系统管理员五种角色，通过自定义装饰器实现视图层的细粒度访问控制。',
        '（2）多级业务审核流程实现。针对不同业务场景实现了差异化的审批机制：由药品管理员审核家庭药箱药品信息；由帖子管理员与药品管理员串行审核共享广场帖子内容及帖子内药品；由用户管理员审核家庭组加入申请。',
        '（3）药品全生命周期管理。构建全局标准药库作为"数据字典"，指导用户规范录入药品信息，实现基于到期日期字段的动态过期预警逻辑，并支持药品管理员对标准药库进行增量同步或覆盖同步。',
        '（4）系统安全与用户体验优化。采用软删除机制保障历史数据安全，利用上下文处理器向全局模板注入消息提醒角标，系统管理员可发布公告与每日医学小贴士，提升平台的信息服务能力。',
    ]
    replace_range(target_start, target_end, goal_texts)
    print('步骤4完成：研究目标替换')

# ══════════════════════════════════════════════════════════
# 第五步：找到"论文结构安排"的提示行并删除
# ══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if '内容：' in para.text and '每章' in para.text:
        para._element.getparent().remove(para._element)
        print(f'删除论文结构提示行 [{i}]')
        break

# ══════════════════════════════════════════════════════════
# 第六步：更新论文结构安排中"第二章"的描述（Vue.js -> Bootstrap 5）
# ══════════════════════════════════════════════════════════
for i, para in enumerate(doc.paragraphs):
    if 'Vue.js' in para.text and 'Django' in para.text and '第2章' in para.text:
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = '第2章　相关技术简介：介绍 Django、Bootstrap 5、MySQL 等核心技术特点及本系统对 B/S 架构的选型依据。'
        else:
            para.add_run('第2章　相关技术简介：介绍 Django、Bootstrap 5、MySQL 等核心技术特点及本系统对 B/S 架构的选型依据。')
        print(f'步骤6完成：更新第2章描述 [{i}]')
        break

# 检查有没有两行拼在一起的（'B/S' 单独一行）
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t == '架构。' or t == '架构' or (len(t) < 6 and 'B/S' in t):
        para._element.getparent().remove(para._element)
        print(f'删除碎片行 [{i}]: {t}')
        break

# ══════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n全部完成，文件已保存。')
