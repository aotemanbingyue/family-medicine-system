# -*- coding: utf-8 -*-
"""
1. 替换 Word 中图4-1~图4-7b 的所有图片
2. 删除 4-7a 前多余的图片段落([193])
"""
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
DIAG = r'd:\SoftWare\Code\bysj\diagrams'

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

def find_cap(*keys, start=0):
    for i in range(start, len(doc.paragraphs)):
        if all(k in doc.paragraphs[i].text for k in keys):
            return i
    return -1

def make_img_elem(path, w_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(w_cm))
    e = p._element
    e.getparent().remove(e)
    return e

def replace_img(img_idx, path, w_cm):
    old = doc.paragraphs[img_idx]._element
    parent = old.getparent()
    pos = list(parent).index(old)
    parent.insert(pos, make_img_elem(path, w_cm))
    parent.remove(old)
    log(f'  替换 [{img_idx}]')

def remove_para(idx):
    p = doc.paragraphs[idx]._element
    p.getparent().remove(p)
    log(f'  删除 [{idx}]')

# ── 诊断当前状态
log('=== 当前图片/图题分布 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    img = has_img(p)
    if img or ('图 4-' in t):
        log(f'  [{i}] img={img}  "{t[:60]}"')

# ── 任务列表：(图题关键词, 图片文件名, 宽度cm)
tasks = [
    (['图 4-1', '系统总体架构图'],   'fig4-1_arch.png',         13.5),
    (['图 4-2', '系统功能模块图'],   'fig4-2_module_tree.png',  14.0),
    (['图 4-3', '家庭组模块流程图'], 'fig4-3_family_flow.png',   8.5),
    (['图 4-4', '药品管理模块流程图'],'fig4-4_medicine_flow.png', 8.5),
    (['图 4-5', '串行双审核'],       'fig4-5_post_flow.png',    14.0),
    (['图 4-6', '系统通知模块流程图'],'fig4-6_notify_flow.png',   8.5),
    (['图 4-7a'],                    'fig4-7a_er_core.png',     13.5),
    (['图 4-7b'],                    'fig4-7b_er_msg.png',      13.5),
]

log('\n=== 开始替换 ===')
for keys, fname, wcm in tasks:
    cap_idx = find_cap(*keys)
    if cap_idx < 0:
        log(f'  未找到图题: {keys}')
        continue
    log(f'\n  图题 [{cap_idx}]: "{doc.paragraphs[cap_idx].text[:50]}"')

    # 在图题前后各找一个图片段落
    img_idx = -1
    # 先找图题前面（最多往前5行）
    for i in range(cap_idx-1, max(cap_idx-6,-1), -1):
        if has_img(doc.paragraphs[i]):
            img_idx = i; break
    # 如果前面没找到，找图题后面
    if img_idx < 0:
        for i in range(cap_idx+1, min(cap_idx+6, len(doc.paragraphs))):
            if has_img(doc.paragraphs[i]):
                img_idx = i; break

    if img_idx < 0:
        log(f'  未找到图片段落，跳过')
        continue

    # 检查 img_idx 前面是否还有额外图片段落（紧挨着的，如4-7a前的[192][193]）
    extra_imgs = []
    for i in range(img_idx-1, max(img_idx-4,-1), -1):
        if has_img(doc.paragraphs[i]):
            extra_imgs.append(i)
        else:
            break
    for ei in extra_imgs:
        log(f'  删除多余图片段落 [{ei}]')
        doc.paragraphs[ei]._element.getparent().remove(doc.paragraphs[ei]._element)
        # 索引偏移修正
        img_idx -= 1
        cap_idx -= 1

    replace_img(img_idx, os.path.join(DIAG, fname), wcm)

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存！')
