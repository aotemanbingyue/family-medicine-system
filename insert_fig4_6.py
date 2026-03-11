# -*- coding: utf-8 -*-
"""
在 初稿 (2).docx 中，在「图 4-6 系统通知模块流程图」图题前插入流程图图片
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys, os

doc_path = r'c:\Users\28041\Desktop\初稿 (2).docx'
out_path = r'c:\Users\28041\Desktop\初稿 (2)_已加图4-6.docx'  # 若原文件被占用则保存到此
img_path = r'd:\SoftWare\Code\bysj\diagrams\fig4-6_notify_flow.png'
CAPTION_KEY = '图 4-6'
CAPTION_FULL = '图 4-6 系统通知模块流程图'

def log(s): sys.stdout.buffer.write((s + '\n').encode('utf-8', 'replace'))

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

doc = Document(doc_path)

# 查找图题段落
cap_idx = -1
for i, p in enumerate(doc.paragraphs):
    if CAPTION_KEY in p.text and '系统通知模块流程图' in p.text:
        cap_idx = i
        log(f'找到图题段落: 索引 [{i}] "{p.text.strip()[:50]}"')
        break

if cap_idx < 0:
    log('未找到「图 4-6 系统通知模块流程图」')
    sys.exit(1)

# 检查图题前是否已有图片
for j in range(cap_idx - 1, max(cap_idx - 6, -1), -1):
    if has_img(doc.paragraphs[j]):
        log('图题前已有图片，无需插入')
        sys.exit(0)
    if doc.paragraphs[j].text.strip():
        break

if not os.path.isfile(img_path):
    log(f'图片不存在: {img_path}')
    sys.exit(1)

# 在图题前插入图片段落（居中）
cap_elem = doc.paragraphs[cap_idx]._element
parent = cap_elem.getparent()
pos = list(parent).index(cap_elem)

img_para = doc.add_paragraph()
img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = img_para.add_run()
run.add_picture(img_path, width=Cm(12))

# 将新段落移到图题正上方
img_elem = img_para._element
img_elem.getparent().remove(img_elem)
parent.insert(pos, img_elem)

doc.save(out_path)
log('已在图题上方插入流程图，已保存: ' + out_path)
