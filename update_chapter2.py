# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def make_para(text, ref_para):
    new_p = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    new_r = OxmlElement('w:r')
    ref_runs = ref_para._element.findall(qn('w:r'))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p

# 第二章内容替换：从 [115]（"2.1 软件架构"）到 [161]（目录结构最后一行）
# [114] "第二章  相关技术简介" 标题保留不动
# 替换 [115] 到 [161]

new_ch2_paras = [
    # 2.1
    '2.1  软件架构',
    '在系统的顶层设计阶段，合理选择软件架构对系统的可维护性、可扩展性与用户体验具有决定性影响。目前常见的架构模式主要包括以下三类。',
    'C/S 架构（Client/Server，客户端/服务器）将应用程序分为客户端与服务端两部分，客户端需安装专用软件，优点是响应速度快、可充分利用本地硬件资源；缺点是客户端维护成本高，每次版本升级需同步更新所有终端，跨平台兼容性较差，对普通家庭用户存在较高的安装与使用门槛。',
    'B/S 架构（Browser/Server，浏览器/服务器）以浏览器为统一客户端，用户无需安装额外程序，所有维护与升级工作集中在服务端完成，对用户完全透明。其优点是跨平台、部署简便、维护成本低，且天然支持多用户协作。移动端原生 APP 与微信小程序虽然在设备交互体验上更优，但开发成本较高，小程序也受微信生态制约，扩展性有所限制。',
    '综合考虑本系统需支持家庭成员跨终端协作管理药品、管理员需在 PC 端进行审核维护、后续可能向社区化方向扩展等需求，B/S 架构是最适合的选择。该模式下，家庭成员无论使用 PC、平板还是手机浏览器均可无缝访问，系统升级只需修改服务端代码即可立即生效，极大地降低了运维负担。因此，本系统采用 B/S 架构进行设计与实现。',
    # 2.2
    '2.2  开发技术',
    '基于上一节确定的 B/S 架构，本系统在后端、前端与数据库三个层面分别选用了经过充分验证的主流技术栈，下面依次介绍各项技术的特点与本系统的选型理由。',
    '2.2.1  Python 与 Django 框架',
    'Python 是一种简洁易读、生态极为丰富的高级编程语言，在 Web 开发、数据处理等领域有广泛应用。本系统后端选用 Python 3.x 结合 Django 5.x 框架进行开发。Django 采用 MVT（Model-View-Template）模式，将数据层、逻辑层与表现层解耦：Model 层通过 ORM 将 Python 类自动映射为数据库表，View 层处理请求与业务逻辑，Template 层负责渲染 HTML 页面。本系统在 models.py 中定义了用户、标准药库、家庭药箱、共享帖子等 9 个核心模型，充分利用 ORM 减少了重复代码。Django 还提供了完善的认证体系，本系统通过继承 AbstractUser 扩展 role 字段，实现五种角色的精细化权限管理，并在 decorators.py 中自定义 @role_required 装饰器对每个视图进行角色校验，有效防止非法越权访问。此外，Django 内置的 CSRF 防护、XSS 过滤与 SQL 注入防御机制进一步保障了系统安全性。',
    '2.2.2  Bootstrap 5 与 Vue.js 前端框架',
    '前端采用 Bootstrap 5 作为主要 CSS 框架，结合 Vue.js 3（CDN 轻量接入）实现响应式与动态化界面。Bootstrap 5 提供了栅格系统、组件库与响应式布局支持，使系统在 PC 端与移动端浏览器上均能获得良好的视觉效果，本系统利用其构建了"手机 App 磁贴"风格的首页仪表盘及多色角标等 UI 元素。Vue.js 3 以 CDN 方式轻量集成，用于实现部分动态交互，结合 context_processors.py 在每次请求中注入待处理数量（如待审核帖子数、未读私聊数），实现了导航栏红点角标的动态展示，提升了用户的信息感知体验。',
    '2.2.3  MySQL 数据库',
    'MySQL 8.0 是目前使用最广泛的开源关系型数据库管理系统之一。本系统以 MySQL 作为主数据库（同时保留 SQLite 3 作为轻量备用方案，通过环境变量 USE_SQLITE 切换）。选用 MySQL 的理由如下：其一，本系统药品信息、用户信息及审核关联关系具有强结构性，关系型数据库的外键约束与事务机制能够保证多表操作的原子性，在家庭成员协作管理药箱时尤为重要；其二，数据库连接采用 utf8mb4 字符集，完整支持中文字符存储；其三，Django 通过 pymysql 驱动与 MySQL 无缝对接，模型迁移与查询优化功能均可直接使用，降低了开发复杂度。',
    # 2.3
    '2.3  开发工具',
    '合理的开发工具链能够显著提升开发效率与代码质量。本系统在代码编写、数据库管理与版本控制等环节分别采用了以下工具。',
    'Cursor（基于 VS Code 架构的 AI 辅助代码编辑器）：兼容 VS Code 全部插件生态，通过代码补全、重构建议与智能诊断功能显著提高了编码规范性与效率，Python 插件提供的 Pylint 静态检查与 Django 语法高亮进一步减少了低级错误。',
    'Navicat for MySQL：广泛使用的数据库可视化管理工具。开发过程中通过 Navicat 对 app_familymedicine、app_sharepost 等核心表进行数据巡检与结构验证，直观确认了 Django 迁移生成的表结构与模型设计的一致性，便于快速定位数据层问题。',
    'Git 版本控制：项目开发过程中采用 Git 进行版本管理，通过提交记录追踪每次功能迭代与 Bug 修复，确保代码变更可追溯、可回滚，为多轮功能迭代提供了可靠的版本保障。',
]

# 获取 ref 段落（用 [116] 的格式作模板）
ref = doc.paragraphs[116]

# 确认替换范围：[115] 到 [161]
start_idx = 115
end_idx = 161

elems_to_remove = [doc.paragraphs[i]._element for i in range(start_idx, end_idx + 1)]
parent = elems_to_remove[0].getparent()
insert_pos = list(parent).index(elems_to_remove[0])

for i, text in enumerate(new_ch2_paras):
    parent.insert(insert_pos + i, make_para(text, ref))

for e in elems_to_remove:
    parent.remove(e)

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('第二章替换完成，文件已保存。')

# 验证
doc2 = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n=== 验证第二章 ===')
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if text and i >= 113 and i <= 145:
        print(f'[{i}] {text[:80]}')
