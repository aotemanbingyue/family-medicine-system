"""
生成第二章"相关技术简介"Word文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 全局字体默认 ──────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, name='黑体', size=16, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, name='黑体', size=14, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, name='黑体', size=13, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(text, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, size=12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def bullet(text, level=1):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=12)
    p.paragraph_format.left_indent  = Cm(level * 0.75)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def code_block(lines):
    """用 Courier New 等宽字体模拟树状代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ═══════════════════════════════════════════════════════════
heading1('第二章  相关技术简介')

# ─── 2.1 ───────────────────────────────────────────────────
heading2('2.1  软件架构')

body(
    '在软件系统的顶层设计阶段，合理选择软件架构对系统的可维护性、可扩展性与用户体验具有'
    '决定性影响。目前主流的软件架构模式主要包括 C/S 架构、B/S 架构以及移动端原生'
    '应用（含微信小程序）三类，下面分别进行分析并说明本系统的选型依据。'
)

heading3('（1）C/S 架构（Client/Server，客户端/服务器）')
body(
    'C/S 架构将应用程序分为客户端与服务器两部分，客户端需要在用户终端安装专用软件，'
    '并通过网络与服务器进行数据交互。其优点是响应速度快、离线能力强、可充分利用本地硬件'
    '资源；缺点是客户端维护成本高，每次版本升级均需同步更新所有终端，跨平台兼容性差，'
    '对非技术背景的家庭用户（如老年人）存在较高的安装与使用门槛。'
)

heading3('（2）B/S 架构（Browser/Server，浏览器/服务器）')
body(
    'B/S 架构以浏览器作为统一客户端，用户无需安装任何额外程序，只需通过标准浏览器即可访问'
    '系统全部功能；系统的维护与升级工作集中在服务端完成，对用户完全透明。其优点是部署简便、'
    '跨平台、维护成本低，且天然支持多用户协作；缺点是在弱网环境下体验略逊于原生客户端，'
    '对浏览器兼容性有一定依赖。'
)

heading3('（3）移动端原生应用与微信小程序')
body(
    '原生 APP 可深度调用设备硬件（摄像头、推送通知等），用户体验最佳，但开发周期长、'
    '需分别维护 iOS 与 Android 两套代码，成本较高。微信小程序无需安装，依托微信生态传播'
    '方便，适合轻量级应用；但其运行环境受微信限制，复杂业务的扩展性有所制约。'
)

heading3('（4）本系统架构选型')
body(
    '综合考虑本系统的应用场景——家庭成员跨终端协作管理药品、管理员需在 PC 端进行审核'
    '与维护、后续可能向社区化方向扩展——B/S 架构是最适合的选择。B/S 模式下，家庭成员'
    '无论使用 PC、平板还是手机浏览器均可无缝访问；系统升级与功能迭代只需修改服务端代码，'
    '所有用户立即生效，极大地降低了运维负担。因此，本系统采用 B/S 架构进行设计与实现。'
)

# ─── 2.2 ───────────────────────────────────────────────────
heading2('2.2  开发技术')

body(
    '基于上一节确定的 B/S 架构，本系统在后端、前端与数据库三个层面分别选用了经过充分'
    '验证的主流技术栈，下面依次介绍各项技术的特点与本系统的选型理由。'
)

# 2.2.1
heading3('2.2.1  Python 与 Django 框架')

body(
    'Python 是一种简洁、易读且生态极为丰富的高级编程语言，在 Web 开发、数据处理与'
    '人工智能等领域均有广泛应用。本系统后端选用 Python（版本 3.x）结合 Django 5.x'
    '框架进行开发。Django 是目前最成熟的 Python Web 框架之一，其核心设计理念是'
    '"不要重复造轮子"（Don\'t Repeat Yourself），具备以下突出优势。'
)

body('一、MVT 设计模式', indent=False)
body(
    'Django 采用 MVT（Model-View-Template）模式，将数据层、逻辑层与表现层解耦。'
    'Model 层通过 ORM（对象关系映射）将 Python 类自动映射为数据库表，开发者无需手写'
    'SQL 即可完成建表与查询；View 层负责处理请求与业务逻辑；Template 层负责渲染'
    'HTML 页面。本系统在 models.py 中定义了用户、标准药库、家庭药箱、共享帖子等'
    '9 个核心模型，充分利用 ORM 减少了重复代码。'
)

body('二、内置认证与权限扩展', indent=False)
body(
    'Django 提供了完善的认证体系。本系统通过继承 AbstractUser 对用户模型进行扩展，'
    '新增 role（角色）字段，实现了普通用户、用户管理员、药品管理员、帖子管理员、'
    '系统管理员五种角色的精细化权限管理。在 decorators.py 中自定义了 @role_required'
    '装饰器，对每个视图进行角色校验，有效防止非法越权访问，保障系统安全性。'
)

body('三、迁移与管理命令', indent=False)
body(
    'Django 内置的数据库迁移机制（makemigrations / migrate）使数据库结构变更可追溯、'
    '可回滚。本系统共完成 9 次迁移，确保模型演进过程有据可查。此外，系统还编写了多个'
    '管理命令（如 create_test_users、seed_demo_scenarios、sync_global_medicines、'
    'auto_publish_medical_tip），用于测试数据生成与自动化运维，提升了开发与演示效率。'
)

body('四、安全性机制', indent=False)
body(
    'Django 默认开启 CSRF 防护、XSS 过滤与 SQL 注入防御，结合系统自定义的角色校验'
    '装饰器与 POST 方法强制校验（@require_POST），形成了多层安全防线。密码存储均经过'
    'Django 标准哈希处理，符合安全开发规范。'
)

# 2.2.2
heading3('2.2.2  Bootstrap 5 与 Vue.js 前端框架')

body(
    '前端采用 Bootstrap 5 作为主要 CSS 框架，结合 Vue.js 3（CDN 轻量接入）实现'
    '响应式与动态化界面。'
)

body('Bootstrap 5', indent=False)
body(
    'Bootstrap 5 是目前使用最广泛的前端 UI 框架，提供了栅格系统、组件库与响应式'
    '布局支持，使系统在 PC 端与移动端浏览器上均能获得良好的视觉效果。本系统利用'
    'Bootstrap 5 构建了"手机 App 磁贴"风格的首页仪表盘、响应式导航栏、卡片式'
    '列表与多色角标等 UI 元素，兼顾了美观性与功能性。'
)

body('Vue.js 3', indent=False)
body(
    'Vue.js 是一款轻量级的渐进式 JavaScript 框架，支持响应式数据绑定与组件化开发。'
    '本系统以 CDN 方式轻量集成 Vue 3，用于实现部分动态交互（如表单联动、实时状态'
    '更新）。结合 context_processors.py 在每次请求中注入待处理数量（如待审核帖子数、'
    '未读私聊数），实现了导航栏红点角标的动态展示，显著提升了用户的信息感知体验。'
)

# 2.2.3
heading3('2.2.3  MySQL 数据库')

body(
    'MySQL 8.0 是世界上使用最广泛的开源关系型数据库管理系统之一。本系统以 MySQL 作为'
    '主数据库（同时保留 SQLite 3 作为轻量备用方案，通过环境变量 USE_SQLITE 切换），'
    '理由如下。'
)

body(
    '第一，关系型结构契合业务需求。本系统的药品信息（生产日期、过期日期）、用户信息与'
    '审核关联关系具有强结构性，关系型数据库的外键约束与事务机制能够保证多表操作的'
    '原子性与一致性，在家庭成员协作管理药箱时尤为重要。'
)

body(
    '第二，字符集与中文支持完善。数据库连接配置采用 utf8mb4 字符集，完整支持中文、'
    '表情符号等字符，避免了汉字存储乱码问题。'
)

body(
    '第三，与 Django ORM 深度兼容。Django 通过 mysqlclient（或 pymysql）驱动与'
    'MySQL 无缝对接，模型迁移、查询优化（select_related、filter 条件索引）等功能'
    '均可直接使用，降低了开发复杂度。'
)

# ─── 2.3 ───────────────────────────────────────────────────
heading2('2.3  开发工具')

body(
    '合理的开发工具链能够显著提升开发效率与代码质量。本系统在代码编写、数据库管理'
    '与项目组织等环节分别采用了以下工具。'
)

heading3('（1）Cursor / Visual Studio Code')
body(
    'Cursor 是基于 VS Code 架构的新一代 AI 辅助代码编辑器，兼容 VS Code 全部插件'
    '生态。本系统主要在 Cursor 中进行开发，借助其代码补全、重构建议与智能诊断功能，'
    '显著提高了编码规范性与效率。VS Code Python 插件提供了 Pylint 静态检查与'
    'Django 语法高亮，进一步减少了低级错误。'
)

heading3('（2）Navicat for MySQL')
body(
    'Navicat 是广泛使用的数据库可视化管理工具。在开发过程中，通过 Navicat 对'
    'app_familymedicine、app_sharepost 等核心表进行数据巡检与结构验证，直观地'
    '确认了 Django 迁移生成的表结构与模型设计的一致性，便于快速定位数据层问题。'
)

heading3('（3）Git 版本控制')
body(
    '项目开发过程中采用 Git 进行版本管理，通过提交记录追踪每次功能迭代与 Bug 修复，'
    '确保代码变更可追溯、可回滚，为多轮功能迭代提供了可靠的版本保障。'
)

# ─── 2.4 ───────────────────────────────────────────────────
heading2('2.4  系统目录结构')

body(
    '本系统采用 Django 标准项目布局，按功能模块划分目录，核心组织结构如下。'
)

tree_lines = [
    'bysj/                          # Django 项目根目录',
    '├── bysj/                      # 项目配置包',
    '│   ├── settings.py            # 全局配置（数据库、中间件、认证等）',
    '│   ├── urls.py                # 根路由分发',
    '│   ├── wsgi.py                # WSGI 部署入口',
    '│   └── asgi.py                # ASGI 部署入口',
    '├── app/                       # 核心业务应用',
    '│   ├── models.py              # 数据模型（9 个实体）',
    '│   ├── views.py               # 视图层（业务逻辑处理）',
    '│   ├── urls.py                # 应用路由（所有 URL 映射）',
    '│   ├── forms.py               # 表单定义与校验',
    '│   ├── decorators.py          # 自定义装饰器（@role_required）',
    '│   ├── context_processors.py  # 上下文处理器（导航角标注入）',
    '│   ├── medicine_sync.py       # 标准药库同步模块',
    '│   ├── medical_tip_auto.py    # 医学小贴士自动发布模块',
    '│   ├── migrations/            # 数据库迁移文件（共 9 次）',
    '│   ├── templates/app/         # HTML 模板（31 个页面）',
    '│   │   ├── base.html          # 全局导航基础模板',
    '│   │   ├── index.html         # 登录用户首页仪表盘',
    '│   │   ├── index_guest.html   # 未登录欢迎页',
    '│   │   ├── family_medicine_*.html   # 家庭药箱相关页面',
    '│   │   ├── share_*.html       # 共享广场相关页面',
    '│   │   ├── message_*.html     # 消息中心与私聊页面',
    '│   │   ├── user_admin_*.html  # 用户管理员后台页面',
    '│   │   ├── medical_tip_*.html # 医学小贴士页面',
    '│   │   └── ...（其余功能页面）',
    '│   ├── static/app/            # 静态资源',
    '│   │   └── style.css          # 全局自定义样式',
    '│   └── management/commands/   # 自定义管理命令',
    '│       ├── create_test_users.py        # 创建测试账号',
    '│       ├── create_family_and_posts.py  # 创建家庭与帖子数据',
    '│       ├── seed_demo_scenarios.py      # 一键生成演示场景数据',
    '│       ├── sync_global_medicines.py    # 标准药库同步命令',
    '│       └── auto_publish_medical_tip.py # 小贴士自动发布命令',
    '├── manage.py                  # Django 管理入口',
    '└── requirements.txt           # Python 依赖清单',
]
code_block(tree_lines)

body(
    '整体目录结构清晰，业务逻辑、模板、静态资源与管理命令各司其职，便于后续功能扩展'
    '与团队协作维护。模板层共包含 31 个 HTML 页面，均继承自 base.html 基础模板，'
    '通过模板继承机制保证了导航栏、页脚与全局样式的统一性，减少了重复代码量。'
)

# ─── 保存 ───────────────────────────────────────────────────
out_path = r'd:\SoftWare\Code\bysj\第二章_相关技术简介.docx'
doc.save(out_path)
print(f'已生成：{out_path}')
