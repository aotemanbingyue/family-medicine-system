# -*- coding: utf-8 -*-
"""
一次性修复：
1. 将所有"图题在图片上方"的图题移到图片下方
2. 将所有"表题在表格上方"的表题移到表格下方  
3. 所有图题/表题字体不加粗、居中
4. 替换参考文献
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import copy, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

# ══════════════════════════════════════════════════════════
# 工具：读取body子元素序列（段落+表格混合）
# ══════════════════════════════════════════════════════════
def get_body_items():
    """返回 [(elem, 'para'/'img'/'table', para_idx_or_None), ...]"""
    body = doc.element.body
    items = []
    para_idx = 0
    for child in body:
        if child.tag == qn('w:p'):
            p = doc.paragraphs[para_idx]
            typ = 'img' if has_img(p) else 'para'
            items.append((child, typ, para_idx))
            para_idx += 1
        elif child.tag == qn('w:tbl'):
            items.append((child, 'table', None))
    return items

def is_fig_caption(txt):
    return bool(txt) and ('图 ' in txt[:8] or txt[:2] in ['图5','图4','图3','图6','图2','图1']) and any(c.isdigit() for c in txt[:12])

def is_tbl_caption(txt):
    return bool(txt) and ('表 ' in txt[:8] or txt[:2] in ['表4','表5','表6','表3','表2','表1']) and any(c.isdigit() for c in txt[:12])

def set_caption_style(para_elem):
    """设置图题/表题：居中、不加粗、字号10.5pt"""
    from docx.oxml.ns import qn as _qn
    # 确保pPr有居中
    pPr = para_elem.find(_qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr'); para_elem.insert(0, pPr)
    jc = pPr.find(_qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); pPr.append(jc)
    jc.set(_qn('w:val'), 'center')
    # 所有run：去加粗、设字号
    for r in para_elem.findall('.//' + _qn('w:r')):
        rPr = r.find(_qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr'); r.insert(0, rPr)
        # 去除加粗
        for b_tag in [_qn('w:b'), _qn('w:bCs')]:
            b = rPr.find(b_tag)
            if b is not None: rPr.remove(b)
        # 设字号21 = 10.5pt
        for existing in rPr.findall(_qn('w:sz')):
            rPr.remove(existing)
        for existing in rPr.findall(_qn('w:szCs')):
            rPr.remove(existing)
        sz_el = OxmlElement('w:sz'); sz_el.set(_qn('w:val'), '21'); rPr.append(sz_el)
        szCs_el = OxmlElement('w:szCs'); szCs_el.set(_qn('w:val'), '21'); rPr.append(szCs_el)

# ══════════════════════════════════════════════════════════
# 核心函数：移动图题到图片后面
# ══════════════════════════════════════════════════════════
def fix_caption_positions():
    """
    扫描body子元素序列：
    - 若发现 [图题段落] 后跟 [图片段落] → 把图题移到图片后面
    - 若发现 [表题段落] 后跟 [表格] → 把表题移到表格后面
    每次修改后重新扫描（因为移动后索引变化）
    """
    max_passes = 20
    for pass_num in range(max_passes):
        body = doc.element.body
        children = list(body)
        moved = False
        
        for i, child in enumerate(children):
            # 判断当前child是否是图题/表题段落
            if child.tag != qn('w:p'):
                continue
            
            # 找到对应段落文字
            txt = ''.join(t.text or '' for t in child.findall('.//' + qn('w:t'))).strip()
            
            is_fig = is_fig_caption(txt)
            is_tbl = is_tbl_caption(txt)
            if not is_fig and not is_tbl:
                continue
            
            # 往前找（跳过空段落）最近的非空child，搜索范围加大到10
            prev_content = None
            for j in range(i-1, max(i-10,-1), -1):
                c = children[j]
                if c.tag == qn('w:tbl'):
                    prev_content = ('table', j, c)
                    break
                elif c.tag == qn('w:p'):
                    ct = ''.join(t.text or '' for t in c.findall('.//' + qn('w:t'))).strip()
                    img = c.find('.//' + qn('w:drawing')) is not None
                    if img:
                        prev_content = ('img', j, c)
                        break
                    elif ct:
                        prev_content = ('para', j, c)
                        break
                    # 空段落继续往前
            
            # 往后找（跳过空段落）最近的非空child，搜索范围加大到10
            next_content = None
            for j in range(i+1, min(i+10, len(children))):
                c = children[j]
                if c.tag == qn('w:tbl'):
                    next_content = ('table', j, c)
                    break
                elif c.tag == qn('w:p'):
                    ct = ''.join(t.text or '' for t in c.findall('.//' + qn('w:t'))).strip()
                    img = c.find('.//' + qn('w:drawing')) is not None
                    if img:
                        next_content = ('img', j, c)
                        break
                    elif ct:
                        next_content = ('para', j, c)
                        break
            
            # 判断是否需要移动
            need_move = False
            target_after = None  # 移到这个elem的后面
            
            if is_fig:
                # 前面不是图片，后面是图片 → 需要移
                prev_is_img = prev_content and prev_content[0] == 'img'
                next_is_img = next_content and next_content[0] == 'img'
                if not prev_is_img and next_is_img:
                    need_move = True
                    target_after = next_content[2]
            
            if is_tbl:
                prev_is_tbl = prev_content and prev_content[0] == 'table'
                next_is_tbl = next_content and next_content[0] == 'table'
                if not prev_is_tbl and next_is_tbl:
                    need_move = True
                    target_after = next_content[2]
            
            if need_move and target_after is not None:
                # 执行移动：先从当前位置移除，再插入到target_after后面
                body.remove(child)
                target_pos = list(body).index(target_after)
                body.insert(target_pos + 1, child)
                set_caption_style(child)
                log(f'  [pass{pass_num+1}] 移动图题/表题: "{txt[:40]}"')
                moved = True
                break  # 重新扫描
        
        if not moved:
            log(f'  pass{pass_num+1}: 无需移动，完成')
            break
    
    return pass_num + 1

# ══════════════════════════════════════════════════════════
# 统一所有图题/表题样式（居中不加粗）
# ══════════════════════════════════════════════════════════
def fix_all_caption_styles():
    count = 0
    for p in doc.paragraphs:
        txt = p.text.strip()
        if is_fig_caption(txt) or is_tbl_caption(txt):
            set_caption_style(p._element)
            count += 1
    log(f'  统一样式: {count} 个图题/表题')

# ══════════════════════════════════════════════════════════
# 替换参考文献
# ══════════════════════════════════════════════════════════
def fix_references():
    # 找到参考文献段落
    ref_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '参考文献':
            ref_idx = i; break
    if ref_idx < 0:
        log('  未找到参考文献段落'); return
    
    log(f'  参考文献位置: [{ref_idx}]')
    
    # 找到致谢段落
    thanks_idx = -1
    for i in range(ref_idx+1, len(doc.paragraphs)):
        if '致谢' in doc.paragraphs[i].text:
            thanks_idx = i; break
    
    log(f'  致谢位置: [{thanks_idx}]')
    
    # 删除参考文献和致谢之间的所有段落（旧文献列表）
    while True:
        cur_thanks = -1
        for i in range(ref_idx+1, len(doc.paragraphs)):
            if '致谢' in doc.paragraphs[i].text:
                cur_thanks = i; break
        if cur_thanks <= ref_idx + 1: break
        doc.paragraphs[ref_idx + 1]._element.getparent().remove(
            doc.paragraphs[ref_idx + 1]._element)
    
    NEW_REFS = [
        '[1] 彭雪,方继莲.AI下乡，撬动基层医疗新变革[N].大众健康报,2025-08-27(005).',
        '[2] 田海晴.基于SpringBoot和Vue框架的共享运营管理平台的设计与实现[D].山东大学,2020.',
        '[3] 劳琳.R医院医共体财务共享平台构建研究[D].广西财经学院,2024.',
        '[4] 张朝阳.关于新时代农村医疗卫生高质量发展内涵、路径与实施的思考[J].中国农村卫生,2025,17(05):12-15.',
        '[5] Westphal A, Mrowka R. Special issue European Journal of Physiology: Artificial intelligence in the field of physiology and medicine[J]. Pflügers Archiv-European Journal of Physiology, 2025:1-4.',
        '[6] Seifieva N E, Zayarnaya A I, Voblaya N I, et al. The organization and management of the system of medical social services in foreign countries[J]. Problemy sotsial\'noi gigieny, zdravookhraneniia i istorii meditsiny, 2024, 32(1):102-105.',
        '[7] 吕冠艳,李奋华.家庭药箱管理平台的开发与实现[J].福建电脑,2023,39(10):83-87.',
    ]
    
    # 找到当前致谢位置
    thanks_now = -1
    for i in range(ref_idx+1, len(doc.paragraphs)):
        if '致谢' in doc.paragraphs[i].text:
            thanks_now = i; break
    
    ref_para = doc.paragraphs[ref_idx]
    anchor_elem = doc.paragraphs[thanks_now]._element
    parent = anchor_elem.getparent()
    pos = list(parent).index(anchor_elem)
    
    for ref_text in reversed(NEW_REFS):
        new_p = OxmlElement('w:p')
        ref_pPr = ref_para._element.find(qn('w:pPr'))
        if ref_pPr is not None:
            new_p.append(copy.deepcopy(ref_pPr))
        new_r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        new_r.append(rPr)
        t = OxmlElement('w:t'); t.text = ref_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(t); new_p.append(new_r)
        parent.insert(pos, new_p)
    
    log(f'  已插入{len(NEW_REFS)}条参考文献')

# ══════════════════════════════════════════════════════════
log('=== Step 1: 移动错误位置的图题/表题 ===')
passes = fix_caption_positions()
log(f'共进行 {passes} 轮扫描')

log('\n=== Step 2: 统一图题/表题样式 ===')
fix_all_caption_styles()

log('\n=== Step 3: 替换参考文献 ===')
fix_references()

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存！')
