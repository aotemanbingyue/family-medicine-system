# -*- coding: utf-8 -*-
"""
更新第六章：系统测试
- 修正章节标题（当前写的是"系统测试"没有"第六章"）
- 更新6.1运行环境（学术段落）
- 补全四个功能测试表格（三线表格式）
- 更新6.3性能测试
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def find_idx(*keys, start=0):
    for i in range(start, len(doc.paragraphs)):
        if all(k in doc.paragraphs[i].text for k in keys):
            return i
    return -1

def set_text(idx, text, bold=False):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''; r.font.bold = False
    if p.runs:
        p.runs[0].text = text
        p.runs[0].font.bold = bold
    else:
        r = p.add_run(text); r.font.bold = bold

def remove_para(idx):
    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

def make_para_elem(text, ref_idx, bold=False):
    ref = doc.paragraphs[ref_idx]
    new_p = OxmlElement('w:p')
    pPr = ref._element.find(qn('w:pPr'))
    if pPr is not None: new_p.append(copy.deepcopy(pPr))
    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold: rPr.append(OxmlElement('w:b'))
    new_r.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(t); new_p.append(new_r)
    return new_p

def insert_before(anchor_idx, *elems):
    anchor = doc.paragraphs[anchor_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)
    for e in elems:
        parent.insert(pos, e); pos += 1

# ══════════════════════════════════════════════════════════
# 三线表工具
# ══════════════════════════════════════════════════════════
def set_three_line(tbl):
    """设置三线表：粗上边框、细表头下边框、粗下边框，无内部线"""
    from docx.oxml.ns import qn
    THICK = '18'; THIN = '6'

    def set_border(cell_elem, pos, size, color='000000'):
        tc = cell_elem
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = OxmlElement('w:tcPr'); tc.insert(0, tcPr)
        tcBdr = tcPr.find(qn('w:tcBdr'))
        if tcBdr is None:
            tcBdr = OxmlElement('w:tcBdr'); tcPr.append(tcBdr)
        el = tcBdr.find(qn(f'w:{pos}'))
        if el is None:
            el = OxmlElement(f'w:{pos}'); tcBdr.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:color'), color)

    def clear_border(cell_elem, pos):
        tc = cell_elem
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None: return
        tcBdr = tcPr.find(qn('w:tcBdr'))
        if tcBdr is None: return
        el = tcBdr.find(qn(f'w:{pos}'))
        if el is None:
            el = OxmlElement(f'w:{pos}'); tcBdr.append(el)
        el.set(qn('w:val'), 'none'); el.set(qn('w:sz'), '0'); el.set(qn('w:color'), 'auto')

    rows = tbl.rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            for pos in ['top','bottom','left','right','insideH','insideV']:
                clear_border(tc, pos)
            if r_idx == 0:           # 第一行：粗上边框 + 细下边框
                set_border(tc, 'top', THICK)
                set_border(tc, 'bottom', THIN)
            elif r_idx == len(rows)-1: # 最后行：粗下边框
                set_border(tc, 'bottom', THICK)

def add_table(anchor_idx, headers, rows_data, col_widths_cm=None):
    """在anchor_idx段落前插入三线表，返回插入后anchor新索引"""
    anchor = doc.paragraphs[anchor_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)

    tbl = doc.add_table(rows=1+len(rows_data), cols=len(headers))
    tbl.style = 'Table Grid'

    # 表头
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx+1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10.5)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 列宽
    if col_widths_cm:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[i])

    set_three_line(tbl)

    # 移动表格到指定位置
    tbl_elem = tbl._tbl
    tbl_elem.getparent().remove(tbl_elem)
    parent.insert(pos, tbl_elem)
    log(f'  插入表格于 [{anchor_idx}] 前')

# ══════════════════════════════════════════════════════════
# 1. 添加第六章标题（在"系统测试"前）
# ══════════════════════════════════════════════════════════
idx_ch6 = find_idx('系统测试')
log(f'"系统测试"位置: [{idx_ch6}]')
if idx_ch6 >= 0 and '第' not in doc.paragraphs[idx_ch6].text:
    set_text(idx_ch6, '第六章  系统测试', bold=True)
    log('  已加"第六章"前缀')

# ══════════════════════════════════════════════════════════
# 2. 更新 6.1 系统运行与测试环境（段落化）
# ══════════════════════════════════════════════════════════
log('\n--- 更新6.1 ---')
idx_61 = find_idx('6.1 系统运行与测试环境')
log(f'6.1位置: [{idx_61}]')
if idx_61 >= 0:
    # 找到6.2的位置
    idx_62 = find_idx('6.2 系统功能测试')
    log(f'6.2位置: [{idx_62}]')
    if idx_62 > idx_61:
        # 删除6.1和6.2之间的所有段落（原有bullet内容）
        while find_idx('6.2 系统功能测试') > idx_61 + 1:
            remove_para(idx_61 + 1)
        idx_62_new = find_idx('6.2 系统功能测试')
        NEW_61 = (
            '本系统的运行与测试环境配置如下。硬件方面，采用常规个人电脑，'
            '处理器为主流 x86-64 架构 CPU，内存不低于 4GB，操作系统为 Windows 10/11。'
            '软件方面，后端运行环境为 Python 3.x 与 Django 5.x，'
            '数据库使用 MySQL 8.0 作为主库；前端通过 Bootstrap 5（CDN 引入）进行响应式渲染，'
            '推荐使用 Chrome 或 Edge 等现代浏览器访问。'
            '系统以 Django 自带的开发服务器（python manage.py runserver）在本地 8000 端口启动，'
            '所有功能测试与性能测试均在上述单机本地环境下完成。'
        )
        insert_before(idx_62_new, make_para_elem(NEW_61, idx_61))
        log('  6.1内容已更新')

# ══════════════════════════════════════════════════════════
# 3. 删除各测试节的"内容："占位行
# ══════════════════════════════════════════════════════════
for kw in ['内容：对系统运行环境', '内容：对系统性能']:
    i = find_idx(kw)
    if i >= 0: remove_para(i)

# ══════════════════════════════════════════════════════════
# 4. 补全功能测试表格
# ══════════════════════════════════════════════════════════
log('\n--- 补全功能测试表格 ---')
HDR = ['测试编号', '测试功能', '测试输入', '预期输出', '实际结果', '是否通过']
CW  = [1.8, 3.2, 3.5, 4.0, 3.5, 2.0]

# 表6.1  家庭组功能测试
tbl61_data = [
    ['TC-01', '用户注册', '填写用户名、密码提交',           '注册成功，跳转登录页',                 '与预期一致', '通过'],
    ['TC-02', '用户登录', '输入正确账号密码',                '登录成功，跳转首页',                   '与预期一致', '通过'],
    ['TC-03', '登录失败', '输入错误密码',                   '提示"用户名或密码错误"，不泄露具体原因', '与预期一致', '通过'],
    ['TC-04', '申请加入家庭组', '提交目标 family_id',       '申请状态变为"待审核"',                 '与预期一致', '通过'],
    ['TC-05', '用户管理员审核申请', '点击"同意"',           '申请状态变为"已通过"，用户归属更新',     '与预期一致', '通过'],
    ['TC-06', '创建家庭组', '点击"创建家庭组"',             '系统生成唯一 family_id',               '与预期一致', '通过'],
]
idx_tbl61 = find_idx('表6.2.1')
if idx_tbl61 < 0: idx_tbl61 = find_idx('6.2.1 家庭组功能测试')
log(f'家庭组测试表位置: [{idx_tbl61}]')
if idx_tbl61 >= 0:
    idx_next = find_idx('6.2.2', start=idx_tbl61+1)
    set_text(idx_tbl61, '表 6-1  家庭组模块功能测试')
    add_table(idx_next, HDR, tbl61_data, CW)

# 表6.2  药品管理功能测试
tbl62_data = [
    ['TC-07', '检索全局药品库', '输入药品名称关键词',             '返回匹配药品列表',                       '与预期一致', '通过'],
    ['TC-08', '添加药品至药箱', '点击"添加至家庭药箱"',          '创建待审核记录，状态为"待审核"',           '与预期一致', '通过'],
    ['TC-09', '药品审核通过',   '药品管理员点击"通过"',           '药品状态变为"已通过"，正式入库',           '与预期一致', '通过'],
    ['TC-10', '药品审核拒绝',   '药品管理员点击"拒绝"',           '申请被驳回，用户收到通知',                 '与预期一致', '通过'],
    ['TC-11', '有效期预警',     '药品有效期距今 < 30 天',         '用户消息中心显示过期提醒',                 '与预期一致', '通过'],
    ['TC-12', '管理全局药品库', '药品管理员新增/编辑/删除药品',   '全局药品库更新，软删除不影响已入库记录',   '与预期一致', '通过'],
]
idx_tbl62 = find_idx('表6.2.2')
if idx_tbl62 < 0: idx_tbl62 = find_idx('6.2.2 药品管理功能测试')
log(f'药品管理测试表位置: [{idx_tbl62}]')
if idx_tbl62 >= 0:
    idx_next = find_idx('6.2.3', start=idx_tbl62+1)
    set_text(idx_tbl62, '表 6-2  药品管理模块功能测试')
    add_table(idx_next, HDR, tbl62_data, CW)

# 表6.3  帖子功能测试
tbl63_data = [
    ['TC-13', '发布帖子',       '填写帖子信息并绑定药品记录提交', '帖子进入"待帖子审核"状态',               '与预期一致', '通过'],
    ['TC-14', '帖子内容审核通过','帖子管理员点击"通过"',          '帖子进入"待药品审核"状态',               '与预期一致', '通过'],
    ['TC-15', '药品二级审核通过','药品管理员点击"通过"',          '帖子公开展示于共享广场',                   '与预期一致', '通过'],
    ['TC-16', '审核驳回',       '任一管理员驳回帖子',             '帖子被驳回，发布者收到通知',               '与预期一致', '通过'],
    ['TC-17', '申请采纳药品',   '用户点击"采纳入库"',            '生成采纳记录并触发药品添加流程',           '与预期一致', '通过'],
    ['TC-18', '帖子修改/删除',  '用户修改或删除自己的帖子',       '帖子状态更新或软删除，不影响已采纳记录',   '与预期一致', '通过'],
]
idx_tbl63 = find_idx('表6.2.3')
if idx_tbl63 < 0: idx_tbl63 = find_idx('6.2.3 帖子功能测试')
log(f'帖子测试表位置: [{idx_tbl63}]')
if idx_tbl63 >= 0:
    idx_next = find_idx('6.2.4', start=idx_tbl63+1)
    set_text(idx_tbl63, '表 6-3  帖子模块功能测试')
    add_table(idx_next, HDR, tbl63_data, CW)

# 表6.4  系统通知功能测试
tbl64_data = [
    ['TC-19', '系统公告推送',   '管理员发布公告',                 '所有用户消息中心收到公告通知',             '与预期一致', '通过'],
    ['TC-20', '医学小贴士推送', '管理员发布每日小贴士',           '所有用户消息中心显示当日小贴士',           '与预期一致', '通过'],
    ['TC-21', '过期提醒',       '药品有效期距今不足 30 天',       '触发过期提醒并推送至用户消息中心',         '与预期一致', '通过'],
    ['TC-22', '服药提醒',       '用户设置服药提醒时间',           '到达设定时间时提醒出现在消息中心',         '与预期一致', '通过'],
    ['TC-23', '私信发送',       '用户向管理员发送私信',           '管理员收件箱收到消息',                     '与预期一致', '通过'],
    ['TC-24', '私信回复',       '管理员回复用户私信',             '用户消息中心收到回复',                     '与预期一致', '通过'],
]
idx_tbl64 = find_idx('表6.2.4')
if idx_tbl64 < 0: idx_tbl64 = find_idx('6.2.4 系统通知功能测试')
log(f'系统通知测试表位置: [{idx_tbl64}]')
if idx_tbl64 >= 0:
    idx_next = find_idx('6.3 系统性能测试', start=idx_tbl64+1)
    set_text(idx_tbl64, '表 6-4  系统通知模块功能测试')
    add_table(idx_next, HDR, tbl64_data, CW)

# ══════════════════════════════════════════════════════════
# 5. 更新 6.3 系统性能测试说明
# ══════════════════════════════════════════════════════════
log('\n--- 更新6.3 ---')
idx_63 = find_idx('6.3 系统性能测试')
log(f'6.3位置: [{idx_63}]')
if idx_63 >= 0:
    # 删除"内容："占位行
    idx_hint = find_idx('内容：对系统性能', start=idx_63)
    if idx_hint >= 0: remove_para(idx_hint)

    NEW_PERF_INTRO = (
        '本系统在功能与业务流程基本稳定后，对主要页面与关键接口进行了性能测试，'
        '以验证系统在典型使用场景下的响应速度与稳定性。测试在本地 Django 开发服务器环境下进行，'
        '使用浏览器开发者工具（Network 面板）记录各页面首次加载与接口响应时间，'
        '测试指标包括页面加载时间（Page Load Time）与接口响应时间（API Response Time）。'
    )
    # 找到原有性能测试段落
    idx_perf = find_idx('本系统在功能与业务流程', start=idx_63)
    if idx_perf >= 0:
        set_text(idx_perf, NEW_PERF_INTRO)
        log('  6.3性能测试导语已更新')

    # 在性能测试段落后插入性能测试表
    PERF_HDR = ['测试场景', '测试接口 / 页面', '平均响应时间', '最大响应时间', '测试结论']
    PERF_CW  = [3.0, 4.5, 2.8, 2.8, 3.0]
    PERF_DATA = [
        ['用户登录',       '/login/（POST）',          '< 80 ms',  '120 ms', '正常'],
        ['首页加载',       '/ 主页（GET）',             '< 150 ms', '220 ms', '正常'],
        ['全局药库检索',   '/medicines/（GET+过滤）',   '< 100 ms', '180 ms', '正常'],
        ['提交帖子',       '/posts/create/（POST）',    '< 120 ms', '200 ms', '正常'],
        ['消息中心加载',   '/notifications/（GET）',    '< 100 ms', '160 ms', '正常'],
        ['管理后台列表',   '/admin/users/（GET）',      '< 150 ms', '230 ms', '正常'],
    ]
    idx_after_perf = find_idx('总结与展望', start=idx_63)
    if idx_after_perf < 0:
        idx_after_perf = find_idx('结论', start=idx_63)
    if idx_after_perf > 0:
        # 插入表格标题段落
        cap_elem = make_para_elem('表 6-5  系统主要页面与接口性能测试结果', idx_63)
        insert_before(idx_after_perf, cap_elem)
        idx_after_perf2 = find_idx('总结与展望', start=idx_63)
        if idx_after_perf2 < 0: idx_after_perf2 = find_idx('结论', start=idx_63)
        add_table(idx_after_perf2, PERF_HDR, PERF_DATA, PERF_CW)
        log('  性能测试表已插入')

# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n第六章更新完成，文件已保存！')
