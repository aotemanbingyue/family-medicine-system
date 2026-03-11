# -*- coding: utf-8 -*-
"""
在 4.4.2 节插入数据库表格，格式与截图一致：
  表标题 + 四列表格（字段名 / 数据类型 / 约束 / 说明）
同时：
  - 删除 4.4.1 和 4.4.2 的"内容："提示行
  - 写入 4.5 安全性设计正文
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

# ── 工具函数 ──────────────────────────────────────────────

def rm_by_content(*keys):
    for i, p in enumerate(doc.paragraphs):
        try:
            t = p.text.strip()
            if all(k in t for k in keys):
                p._element.getparent().remove(p._element)
                print(f'  删除[{i}]: {t[:40]}')
                return True
        except: pass
    return False

def find_idx(*keys):
    for i, p in enumerate(doc.paragraphs):
        try:
            if all(k in p.text for k in keys):
                return i
        except: pass
    return -1

def make_body_para(text, ref_para):
    new_p = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    new_r = OxmlElement('w:r')
    ref_runs = ref_para._element.findall(qn('w:r'))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))
    rPr = new_r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        new_r.insert(0, rPr)
    b = rPr.find(qn('w:b'))
    if b is not None:
        rPr.remove(b)
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p

def set_cell(cell, text, bold=False, center=False, size=10.5):
    cell.text = ''
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table_to_body(body, anchor_elem, table_title, rows_data, ref_para, table_number):
    """
    在 anchor_elem 之前插入：
      1. 空行
      2. 表标题段落
      3. 表格（含表头）
      4. 空行
    """
    parent = anchor_elem.getparent()
    pos = list(parent).index(anchor_elem)

    def ins(elem):
        nonlocal pos
        parent.insert(pos, elem)
        pos += 1

    # 1. 表标题段落
    title_p = make_body_para(f'表 4-{table_number}  {table_title}', ref_para)
    # 让标题居中
    pPr = title_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        title_p.insert(0, pPr)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    # 清除已有 jc
    existing = pPr.find(qn('w:jc'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(jc)
    # 标题加粗
    r = title_p.findall(qn('w:r'))
    for rr in r:
        rPr = rr.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rr.insert(0, rPr)
        b = OxmlElement('w:b')
        rPr.append(b)
    ins(title_p)

    # 2. 表格
    tbl = OxmlElement('w:tbl')
    # 表格属性
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9200')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    tbl.append(tblPr)

    # 列宽定义 (字段名, 类型, 约束, 说明)
    col_widths = [1600, 1600, 2800, 3200]  # twips
    tblGrid = OxmlElement('w:tblGrid')
    for w in col_widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    tbl.append(tblGrid)

    headers = ['字段名', '数据类型', '约束', '说明']
    all_rows = [headers] + rows_data

    for r_idx, row_data in enumerate(all_rows):
        tr = OxmlElement('w:tr')
        is_header = (r_idx == 0)
        for c_idx, cell_text in enumerate(row_data):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(col_widths[c_idx]))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            # 表头底部加粗边框
            if is_header:
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F2F2')
                tcPr.append(shd)
            tc.append(tcPr)
            # 段落
            p = OxmlElement('w:p')
            pPr2 = OxmlElement('w:pPr')
            jc2 = OxmlElement('w:jc')
            jc2.set(qn('w:val'), 'center' if c_idx <= 2 else 'left')
            pPr2.append(jc2)
            p.append(pPr2)
            run = OxmlElement('w:r')
            rPr2 = OxmlElement('w:rPr')
            fn = OxmlElement('w:rFonts')
            fn.set(qn('w:eastAsia'), '宋体')
            fn.set(qn('w:ascii'), '宋体')
            rPr2.append(fn)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '21')  # 10.5pt
            rPr2.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '21')
            rPr2.append(szCs)
            if is_header:
                b2 = OxmlElement('w:b')
                rPr2.append(b2)
            run.append(rPr2)
            t_elem = OxmlElement('w:t')
            t_elem.text = cell_text
            t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            run.append(t_elem)
            p.append(run)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    ins(tbl)

    # 3. 空行
    blank_p = OxmlElement('w:p')
    ins(blank_p)

    print(f'  插入表 4-{table_number}: {table_title}')


# ══════════════════════════════════════════════════════════
# 数据库表定义
# ══════════════════════════════════════════════════════════
TABLES = [
    ('用户表（app_user）', [
        ['id',              'INT',          'PK, AUTO_INCREMENT',     '用户主键'],
        ['username',        'VARCHAR(64)',   'UNIQUE, NOT NULL',       '登录用户名'],
        ['email',           'VARCHAR(128)',  'UNIQUE, NOT NULL',       '邮箱（登录凭证）'],
        ['password',        'VARCHAR(256)',  'NOT NULL',               'Django PBKDF2 加密密码'],
        ['role',            'VARCHAR(20)',   'NOT NULL',               '角色标识（user/admin_user/admin_med/admin_post/admin_sys）'],
        ['family_id',       'VARCHAR(64)',   'NULL',                   '所属家庭组标识，NULL 表示未加入'],
        ['is_active',       'BOOLEAN',       'DEFAULT TRUE',           '账号是否启用'],
        ['date_joined',     'DATETIME',      'DEFAULT NOW',            '注册时间'],
    ]),
    ('全局标准药库表（app_globalmedicine）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '药品主键'],
        ['name',            'VARCHAR(100)',  'UNIQUE, NOT NULL',       '药品通用名称'],
        ['category',        'VARCHAR(50)',   'NOT NULL',               '药品分类（感冒/消炎等）'],
        ['manufacturer',    'VARCHAR(100)',  'NULL',                   '生产厂家'],
        ['description',     'TEXT',          'NULL',                   '药品功效与用法说明'],
        ['is_deleted',      'BOOLEAN',       'DEFAULT FALSE',          '软删除标记'],
        ['created_at',      'DATETIME',      'DEFAULT NOW',            '录入时间'],
    ]),
    ('家庭药箱表（app_familymedicine）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '记录主键'],
        ['user_id',         'INT',           'FK→app_user, NOT NULL',  '所属用户'],
        ['global_med_id',   'INT',           'FK→app_globalmedicine, NULL', '关联标准药库（可为空）'],
        ['name',            'VARCHAR(100)',  'NOT NULL',               '药品名称'],
        ['quantity',        'VARCHAR(50)',   'NULL',                   '库存数量'],
        ['expiration_date', 'DATE',          'NULL',                   '有效期日期'],
        ['audit_status',    'SMALLINT',      'DEFAULT 0',              '审核状态：0待审核 1通过 -1驳回'],
        ['auditor_id',      'INT',           'FK→app_user, NULL',      '审核该条记录的药品管理员'],
        ['reminder_enabled','BOOLEAN',       'DEFAULT FALSE',          '是否开启服药提醒'],
        ['daily_reminder_time','TIME',       'NULL',                   '每日提醒时间'],
        ['dosage_note',     'VARCHAR(200)',  'NULL',                   '剂量备注'],
        ['is_deleted',      'BOOLEAN',       'DEFAULT FALSE',          '软删除标记'],
        ['created_at',      'DATETIME',      'DEFAULT NOW',            '创建时间'],
    ]),
    ('共享帖子表（app_sharepost）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '帖子主键'],
        ['author_id',       'INT',           'FK→app_user, NOT NULL',  '发帖用户'],
        ['family_med_id',   'INT',           'FK→app_familymedicine, NOT NULL', '引用的家庭药箱记录'],
        ['title',           'VARCHAR(200)',  'NOT NULL',               '帖子标题'],
        ['content',         'TEXT',          'NULL',                   '帖子正文'],
        ['status',          'SMALLINT',      'DEFAULT 0',              '帖子审核状态：0待审 1通过 -1驳回'],
        ['auditor_id',      'INT',           'FK→app_user, NULL',      '帖子审核管理员'],
        ['medicine_audit_status','SMALLINT', 'DEFAULT -1',             '药品二级审核：-1待帖审 0待药审 1通过 -2驳回'],
        ['medicine_auditor_id','INT',        'FK→app_user, NULL',      '药品审核管理员'],
        ['is_deleted',      'BOOLEAN',       'DEFAULT FALSE',          '软删除标记'],
        ['created_at',      'DATETIME',      'DEFAULT NOW',            '发布时间'],
    ]),
    ('家庭组加入申请表（app_familyjoinrequest）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '申请主键'],
        ['applicant_id',    'INT',           'FK→app_user, NOT NULL',  '申请加入的用户'],
        ['target_family_id','VARCHAR(64)',   'NOT NULL',               '目标家庭组标识'],
        ['status',          'VARCHAR(10)',   'DEFAULT "pending"',      '状态：pending/approved/rejected'],
        ['reviewer_id',     'INT',           'FK→app_user, NULL',      '审核的用户管理员'],
        ['review_note',     'VARCHAR(200)',  'NULL',                   '审核备注'],
        ['create_time',     'DATETIME',      'DEFAULT NOW',            '申请时间'],
        ['review_time',     'DATETIME',      'NULL',                   '审核完成时间'],
    ]),
    ('私信消息表（app_privatemessage）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '消息主键'],
        ['sender_id',       'INT',           'FK→app_user, NOT NULL',  '发送方用户'],
        ['receiver_id',     'INT',           'FK→app_user, NOT NULL',  '接收方用户'],
        ['content',         'TEXT',          'NOT NULL',               '消息正文'],
        ['create_time',     'DATETIME',      'DEFAULT NOW',            '发送时间'],
        ['read_at',         'DATETIME',      'NULL',                   '已读时间，NULL表示未读'],
    ]),
    ('系统公告表（app_systemannouncement）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '公告主键'],
        ['author_id',       'INT',           'FK→app_user, NOT NULL',  '发布的系统管理员'],
        ['title',           'VARCHAR(200)',  'NOT NULL',               '公告标题'],
        ['content',         'TEXT',          'NOT NULL',               '公告正文'],
        ['created_at',      'DATETIME',      'DEFAULT NOW',            '发布时间'],
    ]),
    ('医学小贴士表（app_medicaltip）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '贴士主键'],
        ['author_id',       'INT',           'FK→app_user, NOT NULL',  '发布的系统管理员'],
        ['tip_date',        'DATE',          'UNIQUE, NOT NULL',       '发布日期（每日唯一）'],
        ['title',           'VARCHAR(200)',  'NOT NULL',               '小贴士标题'],
        ['content',         'TEXT',          'NOT NULL',               '小贴士正文'],
        ['created_at',      'DATETIME',      'DEFAULT NOW',            '创建时间'],
    ]),
    ('帖子药品采纳记录表（app_sharepostmedicineadoption）', [
        ['id',              'INT',           'PK, AUTO_INCREMENT',     '记录主键'],
        ['user_id',         'INT',           'FK→app_user, NOT NULL',  '操作用户'],
        ['post_id',         'INT',           'FK→app_sharepost, NOT NULL', '来源帖子'],
        ['family_med_id',   'INT',           'FK→app_familymedicine, NOT NULL', '新创建的家庭药箱条目'],
        ['created_at',      'DATETIME',      'DEFAULT NOW',            '采纳时间'],
    ]),
]

# ══════════════════════════════════════════════════════════
# 执行：找到插入位置，插入表格
# ══════════════════════════════════════════════════════════

# 删除 4.4.1 的"内容："提示行
rm_by_content('内容：', 'E-R图')
# 删除 4.4.2 的"内容："提示行
rm_by_content('内容：', 'ER图转换')

# 找插入位置：4.5 安全性设计 前面
anchor_idx = find_idx('4.5', '安全性设计')
print(f'插入位置（4.5之前）：[{anchor_idx}]')

anchor_elem = doc.paragraphs[anchor_idx]._element
ref_para = doc.paragraphs[anchor_idx - 1]

# 在 anchor_elem 之前插入引言段落
intro_p = make_body_para(
    '根据系统 ER 图，按照 ER 图转换为关系模式的规则，本系统共设计了 9 张数据库表，具体字段与约束说明如下。',
    ref_para
)
anchor_elem.getparent().insert(
    list(anchor_elem.getparent()).index(anchor_elem),
    intro_p
)

# 重新找 anchor（因为刚插入了一行）
anchor_idx = find_idx('4.5', '安全性设计')
anchor_elem = doc.paragraphs[anchor_idx]._element
body = doc.element.body

for idx, (title, rows) in enumerate(TABLES, 1):
    add_table_to_body(body, anchor_elem, title, rows, ref_para, idx)

# ══════════════════════════════════════════════════════════
# 4.5 安全性设计 - 删除提示行，写入正文
# ══════════════════════════════════════════════════════════
rm_by_content('内容：', '安全性')

anchor_45 = find_idx('系统实现')
anchor_elem_45 = doc.paragraphs[anchor_45]._element
parent_45 = anchor_elem_45.getparent()
pos_45 = list(parent_45).index(anchor_elem_45)

security_texts = [
    '本系统从身份认证、访问控制、数据安全与前端防护四个层面进行了系统性的安全设计。',
    '身份认证与会话管理：系统采用 Django 内置的 Session 认证机制，用户密码使用 PBKDF2 算法进行哈希存储，禁止明文保存。登录失败不提示具体原因（仅返回"用户名或密码错误"），防止用户名枚举攻击。所有需要登录的视图均通过 @login_required 装饰器强制校验，未登录用户自动跳转至登录页。',
    '基于角色的访问控制（RBAC）：系统在 decorators.py 中自定义了 @role_required 装饰器，对每个视图函数进行角色白名单校验。非授权角色访问受保护页面时，系统拒绝请求并重定向至首页，有效防止横向越权与纵向越权操作。',
    'CSRF 防护与 SQL 注入防范：Django 框架默认开启 CSRF 中间件（CsrfViewMiddleware），所有 POST 请求均需携带 CSRF Token 校验，防止跨站请求伪造攻击。系统全程使用 Django ORM 进行数据库操作，ORM 自动对查询参数进行参数化处理，从根本上消除 SQL 注入风险，未使用任何原始 SQL 拼接语句。',
    '数据安全：系统对核心业务数据（药品记录、帖子、公告等）采用软删除（is_deleted 字段）机制，不直接删除数据库记录，保留操作可追溯性。数据库连接配置中的密码通过环境变量（DJANGO_DB_PASSWORD）注入，不硬编码在代码仓库中，避免凭证泄露风险。',
]

for i, text in enumerate(security_texts):
    p = make_body_para(text, ref_para)
    parent_45.insert(pos_45 + i, p)
print(f'插入4.5安全性正文 {len(security_texts)} 段')

# ══════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n全部完成，文件已保存。')
