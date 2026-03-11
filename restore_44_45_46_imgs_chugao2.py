# -*- coding: utf-8 -*-
"""
在《初稿 (2).docx》中，重新为图 4-4、图 4-5、图 4-6 插入流程图图片。
要求：图片所在段落新起一行，段落无缩进、单倍行距、段前后 0 行。
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import os, sys

DOC_PATH = r'c:\Users\28041\Desktop\初稿 (2).docx'
DIAG_DIR = r'd:\SoftWare\Code\bysj\diagrams'

FIGS = [
    ('图 4-4', '药品管理模块流程图', 'fig4-4_medicine_flow.png', 8.5),
    ('图 4-5', '帖子串行双审核流程图', 'fig4-5_post_flow.png', 8.5),
    ('图 4-6', '系统通知模块流程图', 'fig4-6_notify_flow.png', 8.5),
]


def log(s: str):
    sys.stdout.buffer.write((s + '\n').encode('utf-8', 'replace'))


def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None


doc = Document(DOC_PATH)


def find_caption(key, extra):
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if key in t and extra in t:
            return i
    return -1


for key, extra, fname, width_cm in FIGS:
    cap_idx = find_caption(key, extra)
    if cap_idx < 0:
        log(f'未找到图题：{key} {extra}')
        continue

    cap_para = doc.paragraphs[cap_idx]
    log(f'{key} 图题位置: [{cap_idx}] \"{cap_para.text.strip()[:40]}\"')

    # 1. 删除图题附近已存在的图片段落（以防之前脚本插入了不合适的）
    to_del = []
    for i in range(max(0, cap_idx - 5), min(len(doc.paragraphs), cap_idx + 6)):
        if i == cap_idx:
            continue
        if has_img(doc.paragraphs[i]):
            to_del.append(i)
    for idx in sorted(to_del, reverse=True):
        p_elem = doc.paragraphs[idx]._element
        p_elem.getparent().remove(p_elem)
        log(f'  删除旧图片段落 [{idx}]')

    # 2. 插入新段落（在图题上方），设置段落格式，再插入图片
    img_path = os.path.join(DIAG_DIR, fname)
    if not os.path.isfile(img_path):
        log(f'  图片不存在：{img_path}')
        continue

    cap_elem = cap_para._element
    parent = cap_elem.getparent()
    pos = list(parent).index(cap_elem)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.left_indent = Pt(0)
    fmt.right_indent = Pt(0)
    fmt.first_line_indent = Pt(0)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))

    elem = p._element
    elem.getparent().remove(elem)
    parent.insert(pos, elem)

    log(f'  已在图题上方插入图片：{fname}')


doc.save(DOC_PATH)
log('全部处理完成，已保存到原文件。')

