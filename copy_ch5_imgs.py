# -*- coding: utf-8 -*-
"""
从原始初稿复制第五章14张截图，按顺序插入到新文档对应图题前面
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, sys, os, io

orig = Document(r'c:\Users\28041\Desktop\初稿.docx')
doc  = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

# ── 提取原稿中图片字节（按rId）──────────────────────────
def get_img_bytes(src_doc, rId):
    try:
        rel = src_doc.part.rels[rId]
        return rel._target._blob
    except:
        return None

# 原稿第五章图片对应关系（按节顺序）
# 5.1: rId11~14  5.2: rId15~17  5.3: rId18~21  5.4: rId22~24
ch5_imgs = [
    # (rId, 对应图题关键词)
    ('rId11', '图 5-1'),
    ('rId12', '图 5-2'),
    ('rId13', '图 5-3'),
    ('rId14', '图 5-4'),
    ('rId15', '图 5-5'),
    ('rId16', '图 5-6'),
    ('rId17', '图 5-7'),
    ('rId18', '图 5-8'),
    ('rId19', '图 5-9'),
    ('rId20', '图 5-10'),
    ('rId21', '图 5-11'),
    ('rId22', '图 5-12'),
    ('rId23', '图 5-13'),
    ('rId24', '图 5-14'),
]

# ── 先确认新文档当前图片分布 ────────────────────────────
log('=== 新文档第五章当前状态 ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    img = has_img(p)
    if 218 <= i <= 275:
        tag = 'IMG' if img else '   '
        log(f'  [{i}]{tag} {t[:60]}')

# ── 策略：找到每个图题，检查其前面是否有图片，如无则插入 ──
log('\n=== 开始插入图片 ===')

for rId, cap_key in ch5_imgs:
    # 找图题段落
    cap_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if cap_key in p.text and '图 5-' in p.text:
            # 确保是精确匹配（如"图 5-1"不要匹配"图 5-10"）
            txt = p.text.strip()
            # 提取数字
            import re
            num = re.search(r'图 5-(\d+)', txt)
            key_num = re.search(r'图 5-(\d+)', cap_key)
            if num and key_num and num.group(1) == key_num.group(1):
                cap_idx = i
                break

    if cap_idx < 0:
        log(f'  未找到图题: {cap_key}')
        continue

    # 检查图题前面是否已有图片（往前找最多5行）
    already_has = False
    for j in range(cap_idx - 1, max(cap_idx - 6, -1), -1):
        if has_img(doc.paragraphs[j]):
            already_has = True
            break
        if doc.paragraphs[j].text.strip():  # 遇到非空非图片段落
            break

    if already_has:
        log(f'  [{cap_idx}] {cap_key} 已有图片，跳过')
        continue

    # 提取原稿图片字节
    img_bytes = get_img_bytes(orig, rId)
    if not img_bytes:
        log(f'  {rId} 图片字节提取失败')
        continue

    # 在图题段落前插入图片段落
    cap_elem = doc.paragraphs[cap_idx]._element
    parent   = cap_elem.getparent()
    pos      = list(parent).index(cap_elem)

    # 用 python-docx 添加图片段落
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Cm(13.0))

    img_elem = img_para._element
    img_elem.getparent().remove(img_elem)
    parent.insert(pos, img_elem)

    log(f'  [{cap_idx}] 插入图片: {cap_key} (rId={rId}, {len(img_bytes)//1024}KB)')

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n全部完成，文件已保存！')
