# -*- coding: utf-8 -*-
"""
将生成的图片插入 Word 文档对应位置：
- 图4-2 系统功能模块图 → 4.2 节"图 4-2  系统功能模块图"占位行
- 图4-3~4-6 各模块流程图 → 4.2.1~4.2.4 各节文字之后
- 图4-7 ER图 → 4.3.1 实体关系描述之后
"""
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, os

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
DIAG = r'd:\SoftWare\Code\bysj\diagrams'

def find_idx(*keys):
    for i, p in enumerate(doc.paragraphs):
        try:
            if all(k in p.text for k in keys):
                return i
        except: pass
    return -1

def make_para(text, ref_para, bold=False, center=False):
    new_p = OxmlElement('w:p')
    ref_pPr = ref_para._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))
    # 居中
    pPr = new_p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        new_p.insert(0, pPr)
    if center:
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), 'center')
    new_r = OxmlElement('w:r')
    ref_runs = ref_para._element.findall(qn('w:r'))
    if ref_runs:
        ref_rPr = ref_runs[0].find(qn('w:rPr'))
        if ref_rPr is not None:
            new_r.append(copy.deepcopy(ref_rPr))
    rPr = new_r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        new_r.insert(0, rPr)
    b = rPr.find(qn('w:b'))
    if b is not None:
        rPr.remove(b)
    if bold:
        rPr.append(OxmlElement('w:b'))
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    return new_p

def insert_image_before(anchor_idx, img_path, caption, width_cm=13.0, ref_idx=None):
    """在 anchor_idx 段落前插入：图片段 + 图题段"""
    ref = doc.paragraphs[ref_idx if ref_idx is not None else anchor_idx]
    anchor_elem = doc.paragraphs[anchor_idx]._element
    parent = anchor_elem.getparent()
    pos = list(parent).index(anchor_elem)

    # 图片段（居中）
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    img_p_elem = img_para._element
    # 从末尾移除，插入到指定位置
    img_p_elem.getparent().remove(img_p_elem)
    parent.insert(pos, img_p_elem)
    pos += 1

    # 图题段（居中，加粗）
    cap_p = make_para(caption, ref, bold=True, center=True)
    parent.insert(pos, cap_p)

    print(f'  插入图片: {caption}')

def replace_para_text(idx, new_text):
    para = doc.paragraphs[idx]
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

# ══════════════════════════════════════════════════════════
# 1. 图4-2 系统功能模块图 → 找到占位行"图 4-2  系统功能模块图"替换
# ══════════════════════════════════════════════════════════
print('--- 插入图4-2 ---')
idx_42_placeholder = find_idx('图 4-2', '系统功能模块图')
print(f'  占位行: [{idx_42_placeholder}]')
if idx_42_placeholder > 0:
    # 在占位行位置插入图片（替换占位行）
    ref = doc.paragraphs[idx_42_placeholder]
    anchor_elem = ref._element
    parent = anchor_elem.getparent()
    pos = list(parent).index(anchor_elem)

    # 图片段
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(os.path.join(DIAG, 'fig4-2_module_tree.png'), width=Cm(14.0))
    img_p_elem = img_para._element
    img_p_elem.getparent().remove(img_p_elem)
    parent.insert(pos, img_p_elem)

    # 图题替换占位行
    replace_para_text(idx_42_placeholder + 1, '图 4-2  系统功能模块图')
    para = doc.paragraphs[idx_42_placeholder + 1]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.bold = True

# ══════════════════════════════════════════════════════════
# 2. 图4-3~4-6 各模块流程图 → 插在各模块正文段落之后、下一个标题之前
# ══════════════════════════════════════════════════════════
flow_maps = [
    ('4.2.1', '家庭组模块', 'fig4-3_family_flow.png', '图 4-3  家庭组模块流程图'),
    ('4.2.2', '药品管理模块', 'fig4-4_medicine_flow.png', '图 4-4  药品管理模块流程图'),
    ('4.2.3', '帖子模块', 'fig4-5_post_flow.png', '图 4-5  帖子串行双审核流程图'),
    ('4.2.4', '系统通知模块', 'fig4-6_notify_flow.png', '图 4-6  系统通知模块流程图'),
]

print('--- 插入图4-3~4-6 ---')
for sec_num, sec_name, img_file, caption in flow_maps:
    # 找到该节标题的下一个标题位置，在那之前插入图片
    sec_idx = find_idx(sec_num, sec_name)
    if sec_idx < 0:
        print(f'  未找到: {sec_num} {sec_name}')
        continue
    # 找下一个标题（以数字开头的段落）
    next_heading = -1
    for i in range(sec_idx + 1, min(sec_idx + 15, len(doc.paragraphs))):
        t = doc.paragraphs[i].text.strip()
        if t and (t.startswith('4.') or t.startswith('5.')):
            next_heading = i
            break
    if next_heading < 0:
        next_heading = sec_idx + 3

    insert_image_before(
        next_heading,
        os.path.join(DIAG, img_file),
        caption,
        width_cm=9.0 if '流程图' in caption else 13.0,
        ref_idx=sec_idx
    )

# ══════════════════════════════════════════════════════════
# 3. 图4-7 ER图 → 插在 4.3.1 "二、实体间主要关系" 之前
# ══════════════════════════════════════════════════════════
print('--- 插入图4-7 ---')
idx_er = find_idx('二、实体间主要关系')
print(f'  ER插入位置（"二、实体间主要关系"前）: [{idx_er}]')
if idx_er > 0:
    insert_image_before(
        idx_er,
        os.path.join(DIAG, 'fig4-7_er.png'),
        '图 4-7  系统 E-R 实体关系图',
        width_cm=13.5,
        ref_idx=idx_er - 1
    )

# ══════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n全部完成，文件已保存。')
