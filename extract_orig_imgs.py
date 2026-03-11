# -*- coding: utf-8 -*-
"""
从原始初稿提取所有图片，保存到 diagrams/screenshots/ 目录
并列出每张图片在哪个段落（有什么上下文）
"""
from docx import Document
from docx.oxml.ns import qn
import os, sys

orig = Document(r'c:\Users\28041\Desktop\初稿.docx')
OUT = r'd:\SoftWare\Code\bysj\diagrams\screenshots'
os.makedirs(OUT, exist_ok=True)

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

# 提取所有图片（包括关系部分图片）
rels = orig.part.rels
img_rels = {rId: rel for rId, rel in rels.items()
            if hasattr(rel, '_target') and 'image' in str(rel._target).lower()}

log(f'原稿共有图片关系: {len(img_rels)} 个')

saved = {}
for rId, rel in img_rels.items():
    try:
        img_part = rel._target
        img_bytes = img_part._blob
        ext = img_part.partname.split('.')[-1]
        fname = f'{rId}.{ext}'
        fpath = os.path.join(OUT, fname)
        with open(fpath, 'wb') as f:
            f.write(img_bytes)
        saved[rId] = fpath
        log(f'  {rId} -> {fname} ({len(img_bytes)//1024}KB)')
    except Exception as e:
        log(f'  {rId} 提取失败: {e}')

# 找出每张图片在哪个段落
log('\n=== 段落中的图片 ===')
for i, p in enumerate(orig.paragraphs):
    blips = p._element.findall('.//' + qn('a:blip'),
        namespaces={'a':'http://schemas.openxmlformats.org/drawingml/2006/main'})
    if not blips:
        # 尝试另一个命名空间
        blips = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    # 找 r:embed
    for elem in p._element.iter():
        embed = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if embed and embed in saved:
            ctx_before = orig.paragraphs[i-1].text.strip() if i > 0 else ''
            ctx_after  = orig.paragraphs[i+1].text.strip() if i < len(orig.paragraphs)-1 else ''
            log(f'  [{i}] rId={embed}  上文: "{ctx_before[:40]}"  下文: "{ctx_after[:40]}"')
            break

log(f'\n共保存 {len(saved)} 张图片到 {OUT}')
