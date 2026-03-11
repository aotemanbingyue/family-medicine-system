# -*- coding: utf-8 -*-
"""
填写中期检查自查表
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import sys, re, copy

doc = Document(r'd:\SoftWare\Code\bysj\midterm_check.docx')

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def cell_set(cell, text, fontsize=10.5):
    """清空单元格并写入新文字"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = text
            para.runs[0].font.size = Pt(fontsize)
        else:
            run = para.add_run(text)
            run.font.size = Pt(fontsize)
    log(f'  写入: {text[:40]}')

tbl = doc.tables[0]

# 打印行列数
log(f'表格: {len(tbl.rows)}行 x {len(tbl.columns)}列')
for ri, row in enumerate(tbl.rows):
    for ci, cell in enumerate(row.cells):
        t = cell.text.strip()[:30]
        if t:
            log(f'  [{ri},{ci}] {t}')

# ══════════════════════════════════════════════════════════
# 填写各行
# ══════════════════════════════════════════════════════════

# 第0行：学生姓名、专业名称
# 格式：学生姓名 | 学生姓名 | (空) | 专业名称 | 计算机科学与技术
# 需要找到"学生姓名"后的填写格和"专业名称"后的格
row0 = tbl.rows[0]
# 通常第0行结构：[学生姓名][填写值][空][专业名称][计算机科学与技术]
# 先找每格内容
for ci, cell in enumerate(row0.cells):
    if cell.text.strip() in ['学生姓名', '']:
        pass

# 直接按位置填（参考模板结构）
# 行0：学生姓名 _ | 学号 _ | 专业名称 _ | 指导教师 _
# 根据读到的内容，行0有5列：学生姓名 | 学生姓名 | 空 | 专业名称 | 计算机科学与技术

# 找各行
def find_row(*keywords):
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            if any(kw in cell.text for kw in keywords):
                return ri
    return -1

# ── 行0：基本信息 ────────────────────────────────────────
# 结构根据合并情况，col1是姓名值，col4是专业
r0 = tbl.rows[0]
# 找"学生姓名"所在列，下一列就是填写位
for ci in range(len(r0.cells)):
    if r0.cells[ci].text.strip() == '学生姓名':
        cell_set(r0.cells[ci+1], '（你的姓名）')
        break

# 找"专业名称"后的格
for ci in range(len(r0.cells)):
    if '专业名称' in r0.cells[ci].text:
        cell_set(r0.cells[ci+1], '计算机科学与技术')
        break

# ── 行1：任务书、参考文献 ─────────────────────────────────
r1_idx = find_row('任务书')
if r1_idx >= 0:
    r1 = tbl.rows[r1_idx]
    for ci in range(len(r1.cells)):
        if '任务书' in r1.cells[ci].text:
            cell_set(r1.cells[ci+1], '已完成（ √ ），进行中（  ）')
        if '参考文献' in r1.cells[ci].text:
            cell_set(r1.cells[ci+1], '15篇；其中外文文献  2  篇')

# ── 行2：开题报告 ───────────────────────────────────────
r2_idx = find_row('开题报告')
if r2_idx >= 0:
    r2 = tbl.rows[r2_idx]
    for ci in range(len(r2.cells)):
        if '开题报告' in r2.cells[ci].text:
            cell_set(r2.cells[ci+1], '已完成（ √ ），进行中（  ）；完成字数约：3500 字')
            break

# ── 行3：正文 ────────────────────────────────────────────
r3_idx = find_row('正文')
if r3_idx >= 0:
    r3 = tbl.rows[r3_idx]
    for ci in range(len(r3.cells)):
        if r3.cells[ci].text.strip() == '正文':
            cell_set(r3.cells[ci+1], '完成定稿（ √ ），进行中（  ）完成百分比：95%')
            break

# ── 行4：目前已完成任务 ───────────────────────────────────
r4_idx = find_row('目前已完成')
if r4_idx >= 0:
    r4 = tbl.rows[r4_idx]
    for ci in range(len(r4.cells)):
        if '目前已完成' in r4.cells[ci].text:
            NEW_DONE = (
                '1. 需求分析与可行性分析\n'
                '2. 系统总体架构与功能模块设计\n'
                '3. 数据库 E-R 图设计与9张核心表建模\n'
                '4. 基于 Django+Bootstrap 5 的前后端开发\n'
                '5. 五大功能模块全部实现（家庭组、药品、帖子、通知、安全）\n'
                '6. 功能测试与性能测试\n'
                '7. 论文各章节撰写完成'
            )
            cell_set(r4.cells[ci+1], NEW_DONE)
            break

# ── 行5：尚须完成的任务 ───────────────────────────────────
r5_idx = find_row('尚须完成')
if r5_idx >= 0:
    r5 = tbl.rows[r5_idx]
    for ci in range(len(r5.cells)):
        if '尚须完成' in r5.cells[ci].text:
            NEW_TODO = (
                '1. 论文格式与排版最终校对\n'
                '2. 根据指导教师意见修改完善\n'
                '3. 完成论文答辩准备'
            )
            cell_set(r5.cells[ci+1], NEW_TODO)
            break

# ── 行6：存在的问题 ──────────────────────────────────────
r6_idx = find_row('存在的问题')
if r6_idx >= 0:
    r6 = tbl.rows[r6_idx]
    for ci in range(len(r6.cells)):
        if '存在的问题' in r6.cells[ci].text:
            NEW_PROB = (
                '1. 系统目前运行在本地开发环境，尚未进行云端部署与高并发压力测试\n'
                '2. 部分页面 UI 交互细节有待进一步优化\n'
                '3. 论文图表规范与引用格式仍需细化'
            )
            cell_set(r6.cells[ci+1], NEW_PROB)
            break

# ── 行7：采取的办法 ──────────────────────────────────────
r7_idx = find_row('采取的办法')
if r7_idx >= 0:
    r7 = tbl.rows[r7_idx]
    for ci in range(len(r7.cells)):
        if '采取的办法' in r7.cells[ci].text:
            NEW_SOL = (
                '1. 后续计划在本地完成充分的功能验证后，参考 Django 部署文档进行优化\n'
                '2. 与指导教师沟通，按反馈意见逐步修改完善\n'
                '3. 对照论文撰写规范仔细校对格式'
            )
            cell_set(r7.cells[ci+1], NEW_SOL)
            break

# ══════════════════════════════════════════════════════════
doc.save(r'd:\SoftWare\Code\bysj\中期检查自查表_已填写.docx')
log('\n保存完成：中期检查自查表_已填写.docx')
