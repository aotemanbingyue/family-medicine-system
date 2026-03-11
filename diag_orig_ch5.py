# -*- coding: utf-8 -*-
"""对比原始初稿和新文档的第五章图片数量"""
from docx import Document
from docx.oxml.ns import qn
import sys

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

# 原始文档
orig = Document(r'c:\Users\28041\Desktop\初稿.docx')
sys.stdout.buffer.write('=== 原始初稿 第五章图片 ===\n'.encode('utf-8','replace'))
in_ch5 = False
img_count = 0
for i, p in enumerate(orig.paragraphs):
    t = p.text.strip()
    if '系统实现' in t and ('第五章' in t or '5.' in t[:5]):
        in_ch5 = True
    if in_ch5 and ('第六章' in t or '系统测试' in t or '6.1' in t):
        in_ch5 = False
    if in_ch5:
        img = has_img(p)
        if img or t:
            tag = 'IMG' if img else '   '
            sys.stdout.buffer.write(f'[{i}]{tag} {t[:70]}\n'.encode('utf-8','replace'))
            if img: img_count += 1

sys.stdout.buffer.write(f'\n原稿第五章共 {img_count} 张图片\n'.encode('utf-8','replace'))

# 新文档
sys.stdout.buffer.write('\n=== 新文档 第五章图片 ===\n'.encode('utf-8','replace'))
new_doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
in_ch5 = False
img_count2 = 0
for i, p in enumerate(new_doc.paragraphs):
    t = p.text.strip()
    if '第五章' in t and '系统实现' in t:
        in_ch5 = True
    if in_ch5 and ('第六章' in t or '系统测试' in t or '6.1' in t):
        in_ch5 = False
    if in_ch5:
        img = has_img(p)
        if img or (t and ('5.' in t[:4] or '图 5' in t or 'IMG' in t)):
            tag = 'IMG' if img else '   '
            sys.stdout.buffer.write(f'[{i}]{tag} {t[:70]}\n'.encode('utf-8','replace'))
            if img: img_count2 += 1

sys.stdout.buffer.write(f'\n新文档第五章共 {img_count2} 张图片\n'.encode('utf-8','replace'))
