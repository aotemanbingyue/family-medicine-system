# -*- coding: utf-8 -*-
"""
更新第五章：系统实现
- 修正章节标题
- 改写各模块介绍为学术段落
- 清理空行和占位内容
- 删除重复的"安全性实现"节（已在第四章）
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def find_idx(*keys, start=0):
    for i in range(start, len(doc.paragraphs)):
        if all(k in doc.paragraphs[i].text for k in keys):
            return i
    return -1

def set_text(idx, new_text, bold=False):
    para = doc.paragraphs[idx]
    for run in para.runs:
        run.text = ''
        run.font.bold = False
    if para.runs:
        para.runs[0].text = new_text
        para.runs[0].font.bold = bold
    else:
        r = para.add_run(new_text)
        r.font.bold = bold

def remove_para(idx):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

def insert_para_before(anchor_idx, text, bold=False, ref_idx=None):
    """在anchor_idx前插入一个新段落"""
    ref = doc.paragraphs[ref_idx if ref_idx is not None else anchor_idx]
    anchor = doc.paragraphs[anchor_idx]._element
    parent = anchor.getparent()
    pos = list(parent).index(anchor)

    new_p = OxmlElement('w:p')
    ref_pPr = ref._element.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_p.append(copy.deepcopy(ref_pPr))

    new_r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    new_r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(t)
    new_p.append(new_r)
    parent.insert(pos, new_p)

# ══════════════════════════════════════════════════════════
# 1. 修正章节标题"第六章 系统实现" → "第五章 系统实现"
# ══════════════════════════════════════════════════════════
idx = find_idx('第六章 系统实现')
if idx < 0:
    idx = find_idx('第六章', '系统实现')
log(f'章节标题位置: [{idx}]')
if idx >= 0:
    set_text(idx, '第五章  系统实现', bold=True)
    log('  已修正为"第五章  系统实现"')

# ══════════════════════════════════════════════════════════
# 2. 修正 5.4 节下面误用的"5.3.1"标题
# ══════════════════════════════════════════════════════════
# 找到5.4节的"5.3.1 介绍"（错误标题）
idx_wrong = find_idx('5.3.1 介绍', start=310)  # 5.4节之后
log(f'错误标题"5.3.1 介绍"位置: [{idx_wrong}]')
if idx_wrong >= 0:
    set_text(idx_wrong, '5.4.1  介绍', bold=False)

idx_wrong2 = find_idx('5.3.2 实现图', start=320)
log(f'错误标题"5.3.2 实现图"位置: [{idx_wrong2}]')
if idx_wrong2 >= 0:
    set_text(idx_wrong2, '5.4.2  实现图与操作说明', bold=False)

# ══════════════════════════════════════════════════════════
# 3. 改写各模块"介绍"小节 → 学术段落
# ══════════════════════════════════════════════════════════

# --- 5.1 家庭组模块 ---
NEW_51 = (
    '家庭组模块是系统的基础支撑模块，提供用户账号管理与家庭协作关系的建立两项核心功能。'
    '普通用户通过注册界面完成账号创建，填写用户名与密码后账号即可激活；'
    '登录后，用户可通过家庭组 ID 向目标家庭组提交加入申请，由该家庭组的用户管理员进行审批，'
    '审批通过后用户即归属于对应家庭组，可访问家庭药箱等协作功能。'
    '此外，用户可在个人药箱页面对药品库存进行添加与删除操作，'
    '并可通过消息中心接收药品过期提醒，以及通过私信功能与各类管理员进行沟通。'
)

# 找到5.1.1介绍 下面的bullet段落范围，替换为单段
idx_511 = find_idx('5.1.1 介绍')
log(f'5.1.1介绍位置: [{idx_511}]')
if idx_511 >= 0:
    # 找到下一个小标题位置
    next_head = find_idx('5.1.2', start=idx_511+1)
    log(f'5.1.2位置: [{next_head}]')
    # 删除 idx_511+1 到 next_head-1 之间所有段落（旧内容+空行）
    while True:
        # 每次删完后索引变化，重新计算
        cur_next = find_idx('5.1.2', start=idx_511+1)
        if cur_next <= idx_511 + 1:
            break
        remove_para(idx_511 + 1)
    # 在 5.1.2 前插入新段落
    new_next = find_idx('5.1.2', start=idx_511+1)
    insert_para_before(new_next, NEW_51, ref_idx=idx_511)
    log('  5.1 模块介绍已更新')

# --- 5.2 药品管理模块 ---
NEW_52 = (
    '药品管理模块是系统的核心业务模块，涵盖全局标准药库的维护与家庭药箱的全生命周期管理两大功能域。'
    '用户可在全局药品库中按药效分类检索药品，选中后提交"添加至家庭药箱"申请；'
    '该申请将进入待审核状态，由药品管理员对药品信息进行真实性与合规性审核，'
    '审核通过后药品正式入库。药品管理员还负责维护全局标准药库的分类与更新，'
    '并可跨家庭组查阅各用户药箱的库存状态，实现全局监管。'
    '系统在药品有效期临近时（剩余不足 30 天）自动生成过期提醒并推送至用户消息中心。'
)

idx_521 = find_idx('5.2.1 介绍')
log(f'5.2.1介绍位置: [{idx_521}]')
if idx_521 >= 0:
    next_h = find_idx('5.2.2', start=idx_521+1)
    while True:
        cur = find_idx('5.2.2', start=idx_521+1)
        if cur <= idx_521 + 1: break
        remove_para(idx_521 + 1)
    insert_para_before(find_idx('5.2.2', start=idx_521+1), NEW_52, ref_idx=idx_521)
    log('  5.2 模块介绍已更新')

# --- 5.3 帖子模块 ---
NEW_53 = (
    '帖子模块为用户提供药品信息交流与闲置药品转让撮合的功能，采用串行双审核机制保障内容合规。'
    '用户可发布药品转让或求购帖子，帖子需绑定一条家庭药箱中的库存记录以确保药品真实性；'
    '发布后首先由帖子管理员对内容合规性进行审核，通过后再由药品管理员对药品信息进行核实，'
    '两级审核均通过后帖子方可在共享广场公开展示。'
    '其他用户若对帖子中的药品感兴趣，可一键申请将其采纳入自己的家庭药箱；'
    '帖子管理员则可对违规内容进行拦截、删除或处理，并记录处置结果以备审计。'
)

idx_531 = find_idx('5.3.1 介绍')
log(f'5.3.1介绍位置: [{idx_531}]')
if idx_531 >= 0:
    next_h = find_idx('5.3.2', start=idx_531+1)
    while True:
        cur = find_idx('5.3.2', start=idx_531+1)
        if cur <= idx_531 + 1: break
        remove_para(idx_531 + 1)
    insert_para_before(find_idx('5.3.2', start=idx_531+1), NEW_53, ref_idx=idx_531)
    log('  5.3 模块介绍已更新')

# --- 5.4 系统通知模块 ---
NEW_54 = (
    '系统通知模块负责将平台的重要信息主动推送至用户，同时提供用户与管理员之间的私信沟通渠道。'
    '系统在后台定时扫描家庭药箱中各药品的有效期，对剩余有效期不足 30 天的药品自动生成过期提醒；'
    '同时根据用户设置的服药时间生成每日服药提醒，两类提醒均统一推送至用户消息中心。'
    '系统管理员可发布系统公告（如规则更新、功能说明等）及每日医学小贴士，'
    '发布后即时推送至全体用户的消息中心。'
    '用户如需与管理员沟通，可通过私信功能发起会话，管理员在私聊收件箱中查阅并回复，'
    '全部通知与消息统一在用户消息中心集中展示。'
)

idx_541 = find_idx('5.4.1  介绍')
if idx_541 < 0:
    idx_541 = find_idx('5.4', '介绍', start=310)
log(f'5.4.1介绍位置: [{idx_541}]')
if idx_541 >= 0:
    next_h = find_idx('5.4.2', start=idx_541+1)
    if next_h < 0:
        next_h = find_idx('实现图', start=idx_541+1)
    while True:
        cur = find_idx('5.4.2', start=idx_541+1)
        if cur < 0: cur = find_idx('实现图', start=idx_541+1)
        if cur < 0 or cur <= idx_541 + 1: break
        remove_para(idx_541 + 1)
    tgt = find_idx('5.4.2', start=idx_541+1)
    if tgt < 0: tgt = find_idx('实现图', start=idx_541+1)
    if tgt > 0:
        insert_para_before(tgt, NEW_54, ref_idx=idx_541)
    log('  5.4 模块介绍已更新')

# ══════════════════════════════════════════════════════════
# 4. 修正各小标题中的"内容："占位提示和"实现图加上操作说明"
# ══════════════════════════════════════════════════════════
for label in ['内容：功能介绍+实现效果图+操作说明+(可选)关键代码。',
              '内容：功能介绍+实现效果图+操作说明+(可选)关键代码']:
    idx_lbl = find_idx(label)
    while idx_lbl >= 0:
        remove_para(idx_lbl)
        idx_lbl = find_idx(label)

for old, new in [
    ('5.1.2 实现图加上操作说明', '5.1.2  功能实现效果展示'),
    ('5.2.2 实现图加上操作说明', '5.2.2  功能实现效果展示'),
    ('5.3.2 实现图加上操作说明', '5.3.2  功能实现效果展示'),
    ('5.4.2  实现图与操作说明',  '5.4.2  功能实现效果展示'),
]:
    i = find_idx(old)
    if i >= 0:
        set_text(i, new)
        log(f'  标题更新: {new}')

# ══════════════════════════════════════════════════════════
# 5. 删除"5.X 安全性实现"节（内容已在第四章覆盖）
#    替换为指引性文字
# ══════════════════════════════════════════════════════════
idx_sx = find_idx('5.X 安全性实现')
if idx_sx < 0:
    idx_sx = find_idx('安全性实现', start=320)
log(f'5.X安全性节位置: [{idx_sx}]')
if idx_sx >= 0:
    set_text(idx_sx, '5.5  系统安全性实现')
    # 找到下面的说明文字
    idx_sx_desc = idx_sx + 1
    if idx_sx_desc < len(doc.paragraphs):
        NEW_SEC = (
            '系统安全性的具体设计方案已在第四章第 4.4 节中详细阐述，'
            '本节结合系统实现层面，对关键安全措施的落地进行说明。'
            '在身份认证方面，系统使用 Django Session + PBKDF2 哈希机制，'
            '确保用户密码不以明文形式存储；访问控制通过自定义 @role_required 装饰器实现，'
            '对每个视图函数进行角色白名单验证，非授权角色的请求将被重定向至错误页面。'
            '在前端安全方面，所有表单均启用 Django CSRF 中间件保护，'
            '数据库交互全部通过 Django ORM 执行，从根本上规避 SQL 注入风险。'
            '敏感配置（数据库密码、SECRET_KEY 等）通过环境变量注入，不硬编码于版本控制文件中。'
        )
        set_text(idx_sx_desc, NEW_SEC)
        log('  5.5 安全性实现已更新')

# ══════════════════════════════════════════════════════════
doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('\n第五章更新完成，文件已保存！')
