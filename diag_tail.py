# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
import sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
total = len(doc.paragraphs)
sys.stdout.buffer.write(f'总段落数: {total}\n'.encode('utf-8','replace'))
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    img = p._element.find('.//' + qn('w:drawing')) is not None
    if i >= total - 60:
        tag = 'IMG' if img else '   '
        sys.stdout.buffer.write(f'[{i}]{tag} {t[:80]}\n'.encode('utf-8','replace'))
