# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def log(s): sys.stdout.buffer.write((s+'\n').encode('utf-8','replace'))

def set_text(idx, text, bold=False):
    p = doc.paragraphs[idx]
    for r in p.runs: r.text = ''; r.font.bold = False
    if p.runs:
        p.runs[0].text = text; p.runs[0].font.bold = bold
    else:
        r = p.add_run(text); r.font.bold = bold

def remove_para(idx):
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)

def find_idx(*keys, start=0):
    for i in range(start, len(doc.paragraphs)):
        if all(k in doc.paragraphs[i].text for k in keys):
            return i
    return -1

# 末尾总结 [297]
idx_sum = find_idx('总结与展望', start=290)
log(f'末尾总结位置: [{idx_sum}]')
if idx_sum >= 0:
    set_text(idx_sum, '第七章  总结与展望', bold=True)
    log('  已更新')

# 末尾内容查看
total = len(doc.paragraphs)
for i in range(290, total):
    t = doc.paragraphs[i].text.strip()
    log(f'  [{i}] {t[:80]}')

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
log('保存完成')
