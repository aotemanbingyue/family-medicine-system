# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')

def set_para_text(para, new_text):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

# ── 1. 删除 [84] 重复的旧段落（"本研究主要围绕系统的架构设计、功能实现及安全性"旧版本）
#    [84] 是旧的，[85] 是新写的，删掉 [84]
para_84 = doc.paragraphs[84]
para_85 = doc.paragraphs[85]
# 旧的那行含"全面"二字（原文：本研究主要围绕系统的架构设计、功能实现及安全性与可扩展性展开）
# 新的那行含"可扩展性"已被替换，确认下哪个是旧的
print('[84]:', doc.paragraphs[84].text[:60])
print('[85]:', doc.paragraphs[85].text[:60])

# 删掉含"安全性与可扩展性展开"的那行（旧的）
for i, para in enumerate(doc.paragraphs):
    if '安全性与可扩展性展开' in para.text:
        para._element.getparent().remove(para._element)
        print(f'删除重复旧目标行 [{i}]')
        break

# ── 2. 删除论文结构安排的"内容："提示行 [95]
for i, para in enumerate(doc.paragraphs):
    if '本文档的主要内容分为' in para.text or '内容分为以下' in para.text:
        para._element.getparent().remove(para._element)
        print(f'删除论文结构提示行 [{i}]')
        break

# ── 3. 更新第2章描述（[102] 含 Vue.js 的那行 + 紧跟的碎片行）
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if 'Vue.js' in t and 'Django' in t and '第2章' in t:
        set_para_text(para, '第2章　相关技术简介：介绍 Django、Bootstrap 5、MySQL 等核心技术特点及本系统对 B/S 架构的选型依据。')
        print(f'更新第2章描述 [{i}]')
        # 检查下一段是否是碎片（单独的 "架构。" 或 "架构"）
        if i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i + 1]
            nt = next_para.text.strip()
            if nt in ('架构。', '架构', 'B/S', 'B/S 架构。', 'B/S架构。'):
                next_para._element.getparent().remove(next_para._element)
                print(f'删除碎片行 [{i+1}]: {nt}')
        break

doc.save(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
print('\n修复完成，文件已保存。')

# 验证
print('\n=== 验证第一章关键段落 ===')
doc2 = Document(r'c:\Users\28041\Desktop\初稿_新摘要.docx')
for i, para in enumerate(doc2.paragraphs):
    text = para.text.strip()
    if text and i >= 60 and i <= 115:
        print(f'[{i}] {text[:85]}')
