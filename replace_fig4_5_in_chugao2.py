# -*- coding: utf-8 -*-
"""
将桌面上的《初稿 (2).docx》中图 4-5 的帖子双审核流程图替换为竖版（与图4-6风格一致）。
"""
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, sys

DOC_PATH = r'c:\Users\28041\Desktop\初稿 (2).docx'
DIAG_DIR = r'd:\SoftWare\Code\bysj\diagrams'
IMG_NAME = 'fig4-5_post_flow.png'

def log(s):
    sys.stdout.buffer.write((s + '\n').encode('utf-8', 'replace'))

def has_img(p):
    return p._element.find('.//' + qn('w:drawing')) is not None

doc = Document(DOC_PATH)

def find_cap():
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if '图 4-5' in t and '帖子' in t:
            return i
    return -1

cap_idx = find_cap()
if cap_idx < 0:
    log('未找到「图 4-5」的图题段落，脚本结束。')
    sys.exit(1)

log(f'图 4-5 图题位置: [{cap_idx}] "{doc.paragraphs[cap_idx].text.strip()[:50]}"')

# 在图题附近查找现有图片段落
img_idxs = []
for i in range(cap_idx - 6, cap_idx + 7):
    if 0 <= i < len(doc.paragraphs) and has_img(doc.paragraphs[i]):
        img_idxs.append(i)

log(f'附近已有图片段落索引: {img_idxs}')

# 先删除这些旧图片段落（从后往前删，避免索引错乱）
for idx in sorted(img_idxs, reverse=True):
    p = doc.paragraphs[idx]._element
    p.getparent().remove(p)
    log(f'  删除旧图片段落 [{idx}]')

# 在图题上方插入新的竖版图片
img_path = os.path.join(DIAG_DIR, IMG_NAME)
if not os.path.isfile(img_path):
    log('找不到图片文件: ' + img_path)
    sys.exit(1)

cap_elem = doc.paragraphs[cap_idx]._element
parent = cap_elem.getparent()
pos = list(parent).index(cap_elem)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run()
run.add_picture(img_path, width=Cm(8.5))  # 宽度与图4-6一致的竖版效果

elem = para._element
elem.getparent().remove(elem)
parent.insert(pos, elem)

doc.save(DOC_PATH)
log('已替换图 4-5 为竖版流程图，并保存到原文件。')

